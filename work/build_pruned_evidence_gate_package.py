from __future__ import annotations

import csv
import hashlib
import json
import os
import shutil
import subprocess
import sys
import zipfile
from pathlib import Path


ROOT = Path(__file__).resolve().parents[1]
WORK = ROOT / "work"
PROJECT = WORK / "pruned_evidence_gate"
OUTPUTS = ROOT / "outputs"
PACKAGE = OUTPUTS / "Module_14_Capstone_Pruned_Evidence_Gate_Package_Ravi_Bajaj"
PYTHON = Path(r"C:\Users\anaxe\.cache\codex-runtimes\codex-primary-runtime\dependencies\python\python.exe")

for p in (WORK / ".packages", PROJECT):
    if str(p) not in sys.path:
        sys.path.insert(0, str(p))

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from pruned_evidence_gate import CASES, GATE_LABELS, evaluate, write_outputs


BLUE = RGBColor(46, 116, 181)
DARK = RGBColor(31, 77, 120)
INK = RGBColor(0, 0, 0)
MUTED = RGBColor(85, 85, 85)
LIGHT = "F4F6F9"
GRID = "D7DBE2"

SOURCES = [
    ("ML4H 2026 home", "https://ml4h.ahli.cc/"),
    ("ML4H 2026 call for participation", "https://ml4h.ahli.cc/submit/call-for-papers/"),
    (
        "Byrne, D. W., Domenico, H. J., & Moore, R. P. (2024). Artificial intelligence for improved patient outcomes: The pragmatic randomized controlled trial is the secret sauce.",
        "https://kjronline.org/DOIx.php?id=10.3348%2Fkjr.2023.1016",
    ),
    ("Stegenga, J. (2018). Medical Nihilism. Oxford University Press.", "https://global.oup.com/academic/product/medical-nihilism-9780198747048"),
    ("FDA-NIH BEST Resource", "https://www.ncbi.nlm.nih.gov/books/NBK326791/"),
    ("CPIC guidelines overview", "https://cpicpgx.org/guidelines/"),
    ("ClinPGx CPIC CYP2C19-clopidogrel guideline", "https://www.clinpgx.org/guideline/PA166251443"),
    ("CPIC CYP2C19-clopidogrel 2022 update", "https://pmc.ncbi.nlm.nih.gov/articles/PMC9287492/"),
    ("CPIC DPYD-fluoropyrimidines guideline", "https://www.clinpgx.org/guideline/PA166251462"),
    ("CPIC DPYD-fluoropyrimidines update", "https://pmc.ncbi.nlm.nih.gov/articles/PMC5760397/"),
    ("CPIC TPMT/NUDT15-thiopurines update", "https://pmc.ncbi.nlm.nih.gov/articles/PMC12997511/"),
    ("ClinPGx HLA-B-abacavir guideline", "https://www.clinpgx.org/guideline/PA166251444"),
    ("ClinPGx HLA-A/HLA-B-carbamazepine guideline", "https://www.clinpgx.org/guideline/PA166251448"),
    ("ClinVar", "https://www.ncbi.nlm.nih.gov/clinvar/"),
    ("gnomAD", "https://gnomad.broadinstitute.org/"),
    ("All of Us Research Program", "https://allofus.nih.gov/"),
    ("Genome India project", "https://genomeindia.in/"),
    ("Psifas / Mosaic", "https://partnership.psifas.org.il/"),
]


def set_run_font(run, size: float | None = None, color: RGBColor | None = None, bold: bool | None = None, italic: bool | None = None) -> None:
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False, fill: str | None = None) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    set_run_font(run, size=9.5, color=INK, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if fill:
        set_cell_shading(cell, fill)


def set_table_geometry(table, widths: list[int]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_grid = tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        tbl.append(tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for w in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(w))
        tbl_grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Pt(widths[idx] / 20)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(widths[idx]))


def style_doc(doc: Document, footer_label: str) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.45)
    section.footer_distance = Inches(0.35)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.12
    for name, size, color in (("Heading 1", 16, BLUE), ("Heading 2", 13, BLUE), ("Heading 3", 11.5, DARK)):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(4)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run(footer_label)
    set_run_font(run, size=8.5, color=MUTED)


def add_title(doc: Document, title: str, subtitle: str, meta: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    set_run_font(r, size=19, color=DARK, bold=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(subtitle)
    set_run_font(r, size=11.5, color=MUTED, italic=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run(meta)
    set_run_font(r, size=9.5, color=MUTED)


def add_para(doc: Document, text: str, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        set_run_font(r, bold=True)
        r = p.add_run(text[len(bold_prefix):])
        set_run_font(r)
    else:
        r = p.add_run(text)
        set_run_font(r)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(item)
        set_run_font(r)


def add_small_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True, fill=LIGHT)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            set_cell_text(cells[i], val)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_source_list(doc: Document) -> None:
    for label, url in SOURCES:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(f"{label}. {url}")
        set_run_font(r, size=9.2)


def make_figures(summary: dict[str, object], package_figures: Path) -> None:
    import matplotlib.pyplot as plt
    from matplotlib.patches import FancyBboxPatch

    package_figures.mkdir(parents=True, exist_ok=True)

    fig, ax = plt.subplots(figsize=(9, 4.8))
    ax.axis("off")
    boxes = [
        (0.04, 0.62, 0.22, 0.2, "LLM draft alert\nPGx/genomic medication claim"),
        (0.38, 0.62, 0.24, 0.2, "Pre-display evidence gate\n3 checks only"),
        (0.74, 0.62, 0.22, 0.2, "Permitted alert text\nallow, narrow, abstain, deny"),
        (0.16, 0.18, 0.2, 0.18, "Endpoint/actionability"),
        (0.40, 0.18, 0.2, 0.18, "Population fit"),
        (0.64, 0.18, 0.2, 0.18, "Citation/guideline support"),
    ]
    for x, y, w, h, label in boxes:
        ax.add_patch(FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.02", edgecolor="#1F4D78", facecolor="#F4F6F9", linewidth=1.6))
        ax.text(x + w / 2, y + h / 2, label, ha="center", va="center", fontsize=10)
    for start, end in [((0.26, 0.72), (0.38, 0.72)), ((0.62, 0.72), (0.74, 0.72)), ((0.50, 0.62), (0.50, 0.36))]:
        ax.annotate("", xy=end, xytext=start, arrowprops=dict(arrowstyle="->", lw=1.6, color="#333333"))
    ax.text(0.5, 0.04, "Figure 1. Pruned evidence gate: one workflow and three checks.", ha="center", fontsize=9, color="#555555")
    fig.tight_layout()
    fig.savefig(package_figures / "pruned_gate_architecture.png", dpi=180)
    plt.close(fig)

    fig, ax = plt.subplots(figsize=(7.5, 4.8))
    labels = ["Ungated\noverclaim", "Gated remaining\noverclaim", "Sensitivity", "Specificity", "Inappropriate\ndenial"]
    vals = [
        summary["ungated_overclaim_rate"],
        summary["gated_remaining_overclaim_rate"],
        summary["sensitivity_overclaim_detection"],
        summary["specificity_aligned_claim_allowance"],
        summary["inappropriate_denial_count"] / summary["case_count"],
    ]
    colors = ["#C44E52", "#55A868", "#4C72B0", "#4C72B0", "#8172B2"]
    ax.bar(labels, vals, color=colors)
    ax.set_ylim(0, 1.05)
    ax.set_ylabel("Rate")
    ax.set_title("Stage 1 synthetic construct-validity metrics")
    ax.grid(axis="y", alpha=0.25)
    for i, v in enumerate(vals):
        ax.text(i, v + 0.03, f"{v:.2f}", ha="center", fontsize=9)
    fig.tight_layout()
    fig.savefig(package_figures / "pruned_metrics.png", dpi=180)
    plt.close(fig)

    fig, ax = plt.subplots(figsize=(8.2, 4.8))
    gate_counts = dict(summary["primary_gate_counts"])
    labels = [k.replace("_", "\n") for k in gate_counts]
    vals = list(gate_counts.values())
    ax.bar(labels, vals, color="#4C72B0")
    ax.set_ylabel("Cases")
    ax.set_title("Primary gate selected across 30 synthetic cases")
    ax.grid(axis="y", alpha=0.25)
    for i, v in enumerate(vals):
        ax.text(i, v + 0.2, str(v), ha="center", fontsize=9)
    fig.tight_layout()
    fig.savefig(package_figures / "pruned_gate_counts.png", dpi=180)
    plt.close(fig)


def build_proposal(path: Path, summary: dict[str, object]) -> None:
    doc = Document()
    style_doc(doc, "Ravi Bajaj - Pruned Evidence Gate Proposal - August 8, 2026")
    add_title(
        doc,
        "Research Proposal",
        "Evidence-Gated Medical LLM Alerts for Pharmacogenomic Claims",
        "Author: Ravi Bajaj | Target venue: ML4H 2026 | Course: AI in Healthcare",
    )
    add_para(
        doc,
        "Thesis: The capstone is now deliberately pruned. It studies one clinically interpretable workflow: LLM-drafted pharmacogenomic or genomic medication-alert text. The evaluation asks whether a pre-display evidence gate reduces overclaiming relative to guideline-supported evidence without inappropriately denying bounded, aligned alerts.",
        "Thesis:",
    )
    doc.add_heading("Problem Statement", level=1)
    add_para(
        doc,
        "Medical LLMs can sound authoritative even when the evidence only supports a narrower statement. In pharmacogenomics this matters because a real guideline, a verified citation, a population-frequency result, or a variant assertion may support a medication-safety note without proving broad patient-outcome benefit, universal population transport, or deterministic action."
    )
    doc.add_heading("Hypothesis", level=1)
    add_para(
        doc,
        "Compared with ungated LLM alert text, a three-gate controller will reduce evidence overclaiming while tracking inappropriate denial. The supported claim is construct validity, not clinical utility: Stage 1 tests whether the gate enforces the intended evidence grammar on synthetic cases."
    )
    doc.add_heading("Methods", level=1)
    add_bullets(
        doc,
        [
            "Workflow: pharmacogenomic/genomic medication-alert text, including medication-gene pairs and variant assertions relevant to prescribing or safety review.",
            "Gates: endpoint/actionability, population fit, and citation/guideline support.",
            "Stage 1: 30 synthetic cases with expected decisions; executable controller reports overclaim detection, specificity for aligned alerts, inappropriate denial, Brier score, and error categories.",
            "Stage 2: replace author-designed cases with independently authored cases and blinded reviewer adjudication by qualified clinicians or pharmacogenomics reviewers.",
        ],
    )
    doc.add_heading("Stage 1 Progress", level=1)
    add_small_table(
        doc,
        ["Metric", "Current result", "Interpretation"],
        [
            ["Cases", str(summary["case_count"]), "Synthetic case bank for the pruned workflow."],
            ["Ungated overclaim rate", f"{summary['ungated_overclaim_rate']:.2f}", "How often the draft claim is intentionally too strong."],
            ["Gated remaining overclaim", f"{summary['gated_remaining_overclaim_rate']:.2f}", "Construct-validity check, not clinical safety."],
            ["Sensitivity / specificity", f"{summary['sensitivity_overclaim_detection']:.2f} / {summary['specificity_aligned_claim_allowance']:.2f}", "Gate detects overclaims while allowing aligned alerts."],
            ["Inappropriate denial", str(summary["inappropriate_denial_count"]), "Count of bounded aligned alerts blocked too strongly."],
        ],
        [2200, 1700, 5460],
    )
    doc.add_heading("Value", level=1)
    add_para(
        doc,
        "The contribution is a submission-ready narrowing move: instead of presenting a universal assurance theory, the paper proposes a small auditable benchmark for a recurrent clinical workflow. This directly answers the critique that the prior package was broad without a single causal chain."
    )
    doc.add_heading("Selected Sources", level=1)
    add_source_list(doc)
    doc.save(path)


def build_paper(path: Path, summary: dict[str, object], figures: Path) -> None:
    doc = Document()
    style_doc(doc, "Ravi Bajaj - Pruned Evidence Gate Paper Draft - August 8, 2026")
    add_title(
        doc,
        "Evidence-Gated Medical LLM Alerts for Pharmacogenomic Claims",
        "Detecting overclaiming relative to guideline-supported evidence",
        "Author: Ravi Bajaj | Target venue: ML4H 2026 Findings or Proceedings Track | Course: AI in Healthcare",
    )
    doc.add_heading("Abstract", level=1)
    add_para(
        doc,
        "Large language models can draft fluent medication alerts that sound stronger than the evidence they cite. This paper narrows the capstone to one workflow: pharmacogenomic and genomic medication-alert text. I implement a synthetic pre-display evidence gate with three checks: endpoint/actionability, population fit, and citation/guideline support. In 30 synthetic cases spanning CPIC-style medication-gene pairs, variant uncertainty, population transport, and unsupported genetics claims, ungated drafts overclaimed evidence in 20 cases. The gated controller reduced remaining overclaiming to 0/30 while allowing 10/10 aligned bounded alerts and producing no inappropriate denials in this Stage 1 scaffold. These results demonstrate implementation consistency and construct validity against author-designed cases, not independent clinical safety or patient benefit. The planned Stage 2 evaluation compares ungated and gated LLM drafts on independently authored cases with blinded reviewer adjudication of overclaiming, inappropriate denial, calibration, and error categories."
    )
    doc.add_heading("1. Introduction", level=1)
    add_para(
        doc,
        "The prior capstone package showed that a large set of mathematical and assurance ideas could be connected to medical LLM evaluation. The reviewer-facing problem is that breadth weakened the causal chain. This revision makes pruning the method: one workflow, one failure mode, three gates, and a measurable evaluation plan."
    )
    add_para(
        doc,
        "The chosen workflow is pharmacogenomic or genomic medication-alert text. It is attractive because guideline resources such as CPIC and ClinPGx already distinguish genotype, phenotype, evidence level, and prescribing recommendation. That creates a concrete setting where an LLM can be audited for a specific error: converting real but limited evidence into a stronger clinical claim."
    )
    doc.add_heading("2. Related Work and Evidence Boundary", level=1)
    add_para(
        doc,
        "The course theme from Byrne, Domenico, and Moore is that patient-outcome claims for healthcare AI require pragmatic clinical evaluation. Stegenga's Medical Nihilism motivates caution when evidentiary pipelines make interventions appear stronger than they are. FDA-NIH BEST materials provide the endpoint vocabulary: biomarkers, surrogate endpoints, validated surrogates, and clinical outcomes are not interchangeable."
    )
    add_para(
        doc,
        "Pharmacogenomics adds a useful positive control. CPIC-style resources often support bounded prescribing language, such as avoiding a drug in a specific genotype/phenotype context or reviewing dose because of toxicity risk. Those resources do not automatically support claims that an LLM improves outcomes, that an association is causal, that a finding transports across populations without recalibration, or that a weak genetic association authorizes treatment denial."
    )
    doc.add_heading("3. Methods", level=1)
    add_para(
        doc,
        "The controller is a deterministic pre-display gate. It is not presented as an autonomous clinical agent. It receives a drafted alert claim and structured evidence features, then selects one of five permission levels: allow bounded alert, narrow claim, abstain for population fit, abstain for citation/guideline support, or deny unsupported action."
    )
    doc.add_picture(str(figures / "pruned_gate_architecture.png"), width=Inches(6.2))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Figure 1. Pruned architecture for the evidence gate.")
    set_run_font(r, size=9, color=MUTED, italic=True)
    add_small_table(
        doc,
        ["Gate", "Question", "Permitted effect"],
        [
            ["Endpoint/actionability", "Does the evidence support medication action, only surrogate/process language, or no action?", "Allow bounded alert, narrow wording, or deny unsupported action."],
            ["Population fit", "Does source evidence plausibly fit the target patient/cohort?", "Require source-specific validation or abstain from transport."],
            ["Citation/guideline support", "Is the cited guideline real, current, relevant, and strong enough for the claim?", "Deny fabricated claims; abstain on uncertain/conflicting support."],
        ],
        [2100, 3900, 3360],
    )
    doc.add_heading("4. Stage 1 Synthetic Evaluation", level=1)
    add_para(
        doc,
        "The current case bank contains 30 synthetic cases. Ten are bounded aligned alerts; twenty intentionally contain overclaiming through universal action language, unsupported genetics, population mismatch, citation/guideline weakness, variant uncertainty, or obsolete/context-shifted evidence. Each case has an expected action and expected primary gate."
    )
    doc.add_picture(str(figures / "pruned_metrics.png"), width=Inches(5.9))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Figure 2. Stage 1 metrics on author-designed synthetic cases.")
    set_run_font(r, size=9, color=MUTED, italic=True)
    add_small_table(
        doc,
        ["Result", "Value", "Claim allowed"],
        [
            ["Overclaim reduction", f"{summary['absolute_overclaim_reduction']:.2f}", "Construct-validity signal only."],
            ["Sensitivity", f"{summary['sensitivity_overclaim_detection']:.2f}", "Detects intentionally overstrong claims in this synthetic bank."],
            ["Specificity", f"{summary['specificity_aligned_claim_allowance']:.2f}", "Allows aligned bounded alerts in this synthetic bank."],
            ["Brier score", f"{summary['brier_overclaim_risk']:.3f}", "Preliminary risk-score calibration diagnostic."],
            ["Inappropriate denial", str(summary["inappropriate_denial_count"]), "No aligned bounded alerts were denied too strongly in Stage 1."],
        ],
        [2600, 1700, 5060],
    )
    doc.add_picture(str(figures / "pruned_gate_counts.png"), width=Inches(5.9))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Figure 3. Primary gate selected by the deterministic controller.")
    set_run_font(r, size=9, color=MUTED, italic=True)
    doc.add_heading("5. Planned Stage 2", level=1)
    add_para(
        doc,
        "Stage 2 is the part that would make the paper clinically interpretable. It should compare ungated LLM drafts with gated drafts on independently authored cases. Reviewers should be blinded to arm when feasible and should classify overclaiming, inappropriate denial, preserved clinician authority, population-fit concerns, citation support, and error category."
    )
    add_bullets(
        doc,
        [
            "Primary outcome: reduction in unsupported evidence overclaiming.",
            "Safety outcome: inappropriate denial of bounded, guideline-supported alerts.",
            "Secondary outcomes: calibration of overclaim risk, interrater agreement, and qualitative error taxonomy.",
            "Boundary: no real patient data are needed for Stage 1; any Stage 2 use of clinical cases requires appropriate governance or synthetic case authoring.",
        ],
    )
    doc.add_heading("6. Discussion", level=1)
    add_para(
        doc,
        "The key lesson is that an LLM does not need to be proven clinically useful before it can be disciplined. A pre-display gate can enforce a smaller claim grammar: a guideline supports a bounded alert, a population mismatch forces abstention, and an unverifiable citation carries no medical authority. This is not a substitute for pragmatic trials; it is a way to prevent a paper, product, or alert from claiming more than its evidence can bear."
    )
    add_para(
        doc,
        "The most important limitation is that the present results are synthetic and author-designed. They demonstrate that the implementation follows the intended rules, not that the system generalizes. The final paper should be honest about that limitation because it is exactly the epistemic discipline the project asks LLMs to follow."
    )
    doc.add_heading("7. Conclusion", level=1)
    add_para(
        doc,
        "A cohesive capstone can be built from the original implementation by pruning. The revised project studies one practical genetics workflow and evaluates whether a three-gate controller prevents LLM alert text from inflating guideline-supported evidence into unsupported clinical authority."
    )
    doc.add_heading("References and Web Sources", level=1)
    add_source_list(doc)
    doc.save(path)


def build_summary(path: Path, summary: dict[str, object]) -> None:
    doc = Document()
    style_doc(doc, "Ravi Bajaj - Summary Sheet - August 8, 2026")
    add_title(
        doc,
        "Summary Sheet",
        "Pruned evidence-gated pharmacogenomic LLM-alert capstone",
        "Author: Ravi Bajaj | Course: AI in Healthcare",
    )
    doc.add_heading("Conference / Symposium", level=1)
    add_small_table(
        doc,
        ["Field", "Value"],
        [
            ["Name", "Machine Learning for Health Symposium (ML4H 2026)"],
            ["URL", "https://ml4h.ahli.cc/"],
            ["Author instructions / CFP", "https://ml4h.ahli.cc/submit/call-for-papers/"],
            ["Submission deadline", "September 10, 2026, 11:59 PM AoE"],
            ["Author notification", "October 22, 2026"],
            ["Camera-ready deadline", "November 7, 2026 (tentative)"],
            ["Event dates", "December 6-7, 2026"],
            ["Location", "Sydney, Australia"],
            ["Suggested track", "Findings Track if kept as preliminary construct-validity work; Proceedings Track only after independent review."],
        ],
        [2600, 6760],
    )
    doc.add_heading("Title and Abstract", level=1)
    add_para(doc, "Title: Evidence-Gated Medical LLM Alerts for Pharmacogenomic Claims: Detecting Overclaiming Relative to Guideline-Supported Evidence")
    add_para(
        doc,
        "Abstract: This capstone studies a narrowed clinical AI safety problem: LLM-drafted pharmacogenomic and genomic medication alerts can overstate what guideline-supported evidence actually permits. I implement a synthetic pre-display evidence gate with three checks: endpoint/actionability, population fit, and citation/guideline support. In 30 synthetic cases, ungated drafts overclaimed evidence in 20 cases, while the gated controller reduced remaining overclaiming to 0, allowed 10 aligned bounded alerts, and produced no inappropriate denials. These results show implementation consistency and construct validity, not clinical utility. A Stage 2 study would compare ungated and gated alerts on independently authored cases with blinded reviewer adjudication of overclaiming, inappropriate denial, calibration, and error categories."
    )
    doc.add_heading("Generative AI Use", level=1)
    add_para(
        doc,
        "Generative AI tools were used for brainstorming, critique, code scaffolding, drafting, packaging, and formatting. The author selected the final scope, directed the pruning, verified sources, reviewed generated text, ran local tests, and remains responsible for all claims, limitations, and submission materials."
    )
    doc.add_heading("Optional Data, Code, and Tools", level=1)
    add_bullets(
        doc,
        [
            "Synthetic 30-case pharmacogenomic/genomic medication-alert case bank.",
            "Executable three-gate controller and tests.",
            "CSV/JSON results for overclaim detection, inappropriate denial, calibration, and gate/error categories.",
            "No real patient data, no protected health information, no diagnosis, and no treatment recommendation.",
        ],
    )
    doc.save(path)


def build_supplement(path: Path, rows_csv: Path, summary: dict[str, object]) -> None:
    doc = Document()
    style_doc(doc, "Ravi Bajaj - Technical Supplement - August 8, 2026")
    add_title(
        doc,
        "Technical Supplement",
        "Pruned PGx evidence-gate case bank and reproducibility notes",
        "Author: Ravi Bajaj | Course: AI in Healthcare",
    )
    doc.add_heading("Safety Boundary", level=1)
    add_bullets(
        doc,
        [
            "Synthetic-only research artifact.",
            "No real patient data or protected health information.",
            "No diagnosis or treatment recommendation.",
            "The controller evaluates claim permission levels, not patient care.",
            "Stage 1 results are implementation consistency, not independent safety or generalization.",
        ],
    )
    doc.add_heading("Reproducibility", level=1)
    add_para(doc, "Run from the repository root:")
    add_para(doc, r"C:\Users\anaxe\.cache\codex-runtimes\codex-primary-runtime\dependencies\python\python.exe work\pruned_evidence_gate\pruned_evidence_gate.py")
    add_para(doc, r"C:\Users\anaxe\.cache\codex-runtimes\codex-primary-runtime\dependencies\python\python.exe -m pytest work\pruned_evidence_gate -q")
    doc.add_heading("Case Matrix", level=1)
    with rows_csv.open(newline="", encoding="utf-8") as handle:
        reader = csv.DictReader(handle)
        rows = list(reader)
    add_small_table(
        doc,
        ["ID", "Drug-gene", "Action", "Primary gate", "Overclaim"],
        [
            [
                r["case_id"],
                r["drug_gene"],
                r["actual_action"].replace("_", " ").lower(),
                r["primary_gate"].replace("_", " ").lower(),
                r["ungated_overclaim"],
            ]
            for r in rows
        ],
        [900, 1900, 2600, 2500, 1460],
    )
    doc.add_heading("Metric Snapshot", level=1)
    add_small_table(
        doc,
        ["Metric", "Value"],
        [
            ["Case count", str(summary["case_count"])],
            ["Ungated overclaim count", str(summary["ungated_overclaim_count"])],
            ["Gated remaining overclaim count", str(summary["gated_remaining_overclaim_count"])],
            ["Sensitivity", f"{summary['sensitivity_overclaim_detection']:.3f}"],
            ["Specificity", f"{summary['specificity_aligned_claim_allowance']:.3f}"],
            ["Brier score", f"{summary['brier_overclaim_risk']:.3f}"],
            ["Matched expected rate", f"{summary['matched_expected_rate']:.3f}"],
        ],
        [3500, 5860],
    )
    doc.save(path)


def write_readme(path: Path, summary: dict[str, object]) -> None:
    path.write_text(
        f"""# Module 14 Capstone - Pruned Evidence Gate

## Project

Evidence-Gated Medical LLM Alerts for Pharmacogenomic Claims: Detecting Overclaiming Relative to Guideline-Supported Evidence.

## Why this package exists

This package is the pruned, reviewer-aligned version of the capstone. The prior implementation explored many assurance analogies; this submission-facing version narrows to one workflow, three gates, and one measurable failure mode.

## Workflow

LLM-drafted pharmacogenomic or genomic medication-alert text.

## Gates

1. Endpoint/actionability.
2. Population fit.
3. Citation/guideline support.

## Stage 1 result

- Synthetic cases: {summary['case_count']}
- Ungated overclaim rate: {summary['ungated_overclaim_rate']:.2f}
- Gated remaining overclaim rate: {summary['gated_remaining_overclaim_rate']:.2f}
- Sensitivity: {summary['sensitivity_overclaim_detection']:.2f}
- Specificity: {summary['specificity_aligned_claim_allowance']:.2f}
- Inappropriate denial count: {summary['inappropriate_denial_count']}

These are construct-validity results on author-designed synthetic cases. They do not prove clinical safety, accuracy, generalization, or patient benefit.

## Stage 2 plan

Compare ungated and gated LLM alert drafts on independently authored cases with blinded reviewer adjudication. Report overclaim reduction, inappropriate denial, calibration, interrater agreement, and error categories.

## Contents

- `proposal/`: research proposal in DOCX and PDF.
- `paper/`: paper draft in DOCX and PDF.
- `summary/`: required 1-2 page summary sheet in DOCX and PDF.
- `supplement/`: case matrix and reproducibility notes.
- `code/`: pruned evaluator and tests only.
- `results/`: CSV and JSON outputs.
- `figures/`: three submission-facing figures.

## Safety boundary

No real patient data, no protected health information, no diagnosis, and no treatment recommendation.
""",
        encoding="utf-8",
    )


def write_safety(path: Path) -> None:
    path.write_text(
        """# Safety and Data Boundary

This is a synthetic course research artifact. It is not a medical product and does not recommend diagnosis or treatment.

The package uses no real patient records, no protected health information, and no private genomic data. Public guideline and database names are used as evidence anchors and population-fit examples only.

The implemented gate checks whether alert text should be allowed, narrowed, abstained, or denied as a matter of claim permission. Clinician or pharmacist authority remains outside the model.
""",
        encoding="utf-8",
    )


def export_pdf(docx_path: Path, pdf_path: Path) -> int:
    proc = subprocess.run(
        [str(PYTHON), str(WORK / "export_docx_with_word.py"), str(docx_path), str(pdf_path)],
        cwd=ROOT,
        text=True,
        capture_output=True,
        check=True,
    )
    pages = 0
    for line in proc.stdout.splitlines():
        if line.startswith("WordPages="):
            pages = int(line.split("=", 1)[1])
    return pages


def sha256(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def write_manifest(package: Path) -> None:
    rows = []
    for path in sorted(p for p in package.rglob("*") if p.is_file() and p.name != "MANIFEST_SHA256.csv"):
        rows.append((str(path.relative_to(package)).replace("\\", "/"), path.stat().st_size, sha256(path)))
    manifest = package / "MANIFEST_SHA256.csv"
    with manifest.open("w", newline="", encoding="utf-8") as handle:
        writer = csv.writer(handle)
        writer.writerow(["path", "bytes", "sha256"])
        writer.writerows(rows)


def zip_package(package: Path) -> Path:
    zip_path = package.with_suffix(".zip")
    if zip_path.exists():
        zip_path.unlink()
    with zipfile.ZipFile(zip_path, "w", compression=zipfile.ZIP_DEFLATED) as z:
        for path in sorted(p for p in package.rglob("*") if p.is_file()):
            z.write(path, package.name + "/" + str(path.relative_to(package)).replace("\\", "/"))
    return zip_path


def build_package() -> dict[str, object]:
    summary = write_outputs(PROJECT / "results")
    if PACKAGE.exists():
        shutil.rmtree(PACKAGE)
    for sub in ("code", "results", "figures", "proposal", "paper", "summary", "supplement"):
        (PACKAGE / sub).mkdir(parents=True, exist_ok=True)

    shutil.copy2(PROJECT / "pruned_evidence_gate.py", PACKAGE / "code" / "pruned_evidence_gate.py")
    shutil.copy2(PROJECT / "test_pruned_evidence_gate.py", PACKAGE / "code" / "test_pruned_evidence_gate.py")
    for result in (PROJECT / "results").glob("*"):
        shutil.copy2(result, PACKAGE / "results" / result.name)

    make_figures(summary, PROJECT / "figures")
    for fig in (PROJECT / "figures").glob("pruned_*.png"):
        shutil.copy2(fig, PACKAGE / "figures" / fig.name)

    proposal = PACKAGE / "proposal" / "Module_14_Capstone_Proposal_Pruned_Evidence_Gate_Ravi_Bajaj.docx"
    paper = PACKAGE / "paper" / "Module_14_Final_Paper_Pruned_Evidence_Gate_Ravi_Bajaj.docx"
    summary_doc = PACKAGE / "summary" / "Module_14_Summary_Sheet_Pruned_Evidence_Gate_Ravi_Bajaj.docx"
    supplement = PACKAGE / "supplement" / "Module_14_Technical_Supplement_Pruned_Evidence_Gate_Ravi_Bajaj.docx"
    build_proposal(proposal, summary)
    build_paper(paper, summary, PACKAGE / "figures")
    build_summary(summary_doc, summary)
    build_supplement(supplement, PACKAGE / "results" / "pruned_pgx_case_results.csv", summary)

    page_counts = {
        "proposal": export_pdf(proposal, proposal.with_suffix(".pdf")),
        "paper": export_pdf(paper, paper.with_suffix(".pdf")),
        "summary": export_pdf(summary_doc, summary_doc.with_suffix(".pdf")),
        "supplement": export_pdf(supplement, supplement.with_suffix(".pdf")),
    }
    write_readme(PACKAGE / "README.md", summary)
    write_safety(PACKAGE / "SAFETY.md")
    write_manifest(PACKAGE)
    zip_path = zip_package(PACKAGE)
    return {"summary": summary, "page_counts": page_counts, "package": str(PACKAGE), "zip": str(zip_path)}


def main() -> None:
    result = build_package()
    print(json.dumps(result, indent=2))


if __name__ == "__main__":
    main()
