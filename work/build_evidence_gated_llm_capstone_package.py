from __future__ import annotations

import csv
import hashlib
import json
import os
import shutil
import subprocess
import sys
import zipfile
from pathlib import Path


ROOT = Path(__file__).resolve().parents[1]
WORK = ROOT / "work"
PROJECT = WORK / "evidence_gated_llm_capstone"
OUTPUTS = ROOT / "outputs"
PACKAGE = OUTPUTS / "Module_14_Capstone_Evidence_Gated_Medical_LLMs_Package_Ravi_Bajaj"
PYTHON = Path(r"C:\Users\anaxe\.cache\codex-runtimes\codex-primary-runtime\dependencies\python\python.exe")

if str(WORK / ".packages") not in sys.path:
    sys.path.insert(0, str(WORK / ".packages"))

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BLUE = RGBColor(46, 116, 181)
DARK = RGBColor(31, 77, 120)
INK = RGBColor(0, 0, 0)
MUTED = RGBColor(85, 85, 85)
LIGHT = "F4F6F9"
GRID = "D7DBE2"

ACTION_SHORT = {
    "ALLOW_BOUNDED_SECOND_READER": "allow second reader",
    "ALLOW_BOUNDED_TRAINING_SIMULATION": "allow training simulation",
    "ALLOW_CAUTIOUS_CLINICIAN_SUMMARY": "allow outcome summary",
    "DENY_CITATION": "deny citation",
    "ABSTAIN_TRANSPORT": "abstain transport",
    "ESCALATE_WORKFLOW_DIAGNOSIS": "workflow diagnosis",
    "ABSTAIN_OPPORTUNITY_COST": "abstain cost",
    "ABSTAIN_EVIDENCE_CHAIN": "abstain chain",
    "ABSTAIN_CONTEXT": "abstain context",
    "STRESS_TEST_GENERALIZATION": "stress test shift",
    "ABSTAIN_PROVENANCE": "abstain provenance",
    "NARROW_WITH_CONFIRMATION": "confirmatory narrow",
    "NARROW_TO_SURROGATE": "narrow surrogate",
    "STRESS_TEST_ORDERING": "stress test ordering",
    "STRESS_TEST_CONFOUNDING": "stress test confounding",
    "PRESERVE_CONSENT_BOUNDARY": "preserve consent boundary",
}


SOURCES = [
    (
        "Stegenga, J. (2018). Medical Nihilism. Oxford University Press.",
        "https://global.oup.com/academic/product/medical-nihilism-9780198747048",
    ),
    (
        "Eronen, M. I. (2019). Review of Medical Nihilism. Notre Dame Philosophical Reviews.",
        "https://ndpr.nd.edu/reviews/medical-nihilism/",
    ),
    (
        "Byrne, D. W., Domenico, H. J., & Moore, R. P. (2024). Artificial intelligence for improved patient outcomes: The pragmatic randomized controlled trial is the secret sauce. Korean Journal of Radiology, 25(2), 123-125.",
        "https://kjronline.org/DOIx.php?id=10.3348%2Fkjr.2023.1016",
    ),
    (
        "FDA-NIH Biomarker Working Group. (2016-2025). BEST (Biomarkers, EndpointS, and other Tools) Resource.",
        "https://www.ncbi.nlm.nih.gov/books/NBK326791/",
    ),
    (
        "U.S. Food and Drug Administration. FDA facts: Biomarkers and surrogate endpoints.",
        "https://www.fda.gov/about-fda/innovation-fda/fda-facts-biomarkers-and-surrogate-endpoints",
    ),
    (
        "Lee, C. R., Luzum, J. A., Sangkuhl, K., et al. (2022). Clinical Pharmacogenetics Implementation Consortium guideline for CYP2C19 genotype and clopidogrel therapy: 2022 update.",
        "https://pmc.ncbi.nlm.nih.gov/articles/PMC9287492/",
    ),
    (
        "Clinical Genome Resource. Current ACMG Secondary Findings List.",
        "https://search.clinicalgenome.org/kb/genes/acmgsf",
    ),
    (
        "Health Resources and Services Administration. Recommended Uniform Screening Panel.",
        "https://newbornscreening.hrsa.gov/about-newborn-screening/recommended-uniform-screening-panel",
    ),
    (
        "American College of Medical Genetics and Genomics. ACT Sheets and Algorithms.",
        "https://www.acmg.net/ACMG/Medical-Genetics-Practice-Resources/ACT_Sheets_and_Algorithms.aspx",
    ),
    (
        "Centers for Disease Control and Prevention. Tier 1 Genomics Applications and their Importance to Public Health.",
        "https://archive.cdc.gov/www_cdc_gov/genomics/implementation/toolkit/tier1.htm",
    ),
    (
        "Walker, S. C., French, B., Moore, R. P., et al. (2023). Model-guided decision-making for thromboprophylaxis and hospital-acquired thromboembolic events among hospitalized children and adolescents: The CLOT randomized clinical trial. JAMA Network Open, 6(10), e2337789.",
        "https://jamanetwork.com/journals/jamanetworkopen/fullarticle/2810644",
    ),
    (
        "Carroll, G. D. (2005). Formal properties of categorial grammars. Harvard University undergraduate thesis.",
        "https://www.math.harvard.edu/media/gabriel_carroll.pdf",
    ),
    (
        "Kanazawa, M. (1998). Learnable classes of categorial grammars. CSLI Publications.",
        "https://eprints.illc.uva.nl/id/eprint/1971/",
    ),
    (
        "Vishnikin, M., & Okhotin, A. (2025). Categorial grammars with unique category assignment. arXiv:2505.14559.",
        "https://arxiv.org/abs/2505.14559",
    ),
    (
        "Chatterjee, S. (2022). Superconcentration in surface growth. arXiv:2103.09199.",
        "https://arxiv.org/abs/2103.09199",
    ),
    (
        "Yablo, S. (1993). Paradox without self-reference.",
        "https://www.mit.edu/~yablo/pwsr.pdf",
    ),
    (
        "The Story of Mathematics. Paul Cohen: Set theory and the continuum hypothesis.",
        "https://www.storyofmathematics.com/20th_cohen.html/",
    ),
    (
        "ML4H 2026. Call for Participation.",
        "https://ml4h.ahli.cc/submit/call-for-papers/",
    ),
    (
        "Lynn, B. Binary Decision Diagrams: Families and ZDDs. Stanford Crypto Notes.",
        "https://crypto.stanford.edu/pbc/notes/zdd/",
    ),
    (
        "Akaike, H. (1974). A new look at the statistical model identification. IEEE Transactions on Automatic Control, 19(6), 716-723.",
        "https://ieeexplore.ieee.org/document/1100705",
    ),
    (
        "Schwarz, G. (1978). Estimating the dimension of a model. The Annals of Statistics, 6(2), 461-464.",
        "https://projecteuclid.org/journals/annals-of-statistics/volume-6/issue-2/Estimating-the-Dimension-of-a-Model/10.1214/aos/1176344136.short",
    ),
    (
        "Hou, Y., Ji, T., Zhang, D., & Stefanidis, A. (2025). Kolmogorov-Arnold Networks: A critical assessment of claims, performance, and practical viability. arXiv:2407.11075.",
        "https://arxiv.org/abs/2407.11075",
    ),
    (
        "GeneReviews. Genetic Hearing Loss Overview.",
        "https://www.ncbi.nlm.nih.gov/books/NBK1434/",
    ),
    (
        "Cohen, S. M., et al. Functional outcomes and quality of life after cochlear implantation.",
        "https://pmc.ncbi.nlm.nih.gov/articles/PMC9457208/",
    ),
    (
        "Hearing Loss Association of America. 'CODA' movie sheds light on hearing loss.",
        "https://www.hearingloss.org/coda-movie-sheds-light/",
    ),
    (
        "Association of American Medical Colleges. Telehealth Competencies Across the Learning Continuum.",
        "https://store.aamc.org/downloadable/download/sample/sample_id/412/",
    ),
    (
        "Crandall, M. G., & Lions, P. L. (1983). Viscosity solutions of Hamilton-Jacobi equations.",
        "https://www.ams.org/journals/tran/1983-277-01/S0002-9947-1983-0690039-8/",
    ),
    (
        "Korteweg, D. J., & de Vries, G. (1895). On the change of form of long waves advancing in a rectangular canal, and on a new type of long stationary waves.",
        "https://doi.org/10.1080/14786449508620739",
    ),
    (
        "Wolber, L. E., et al. (2014). Salt-inducible kinase 3, SIK3, is a new gene associated with hearing. Human Molecular Genetics, 23(23), 6407-6418.",
        "https://doi.org/10.1093/hmg/ddu346",
    ),
    (
        "Bhatt, I. S., Wilson, N., Dias, R., & Torkamani, A. (2022). A genome-wide association study of tinnitus reveals shared genetic links to neuropsychiatric disorders. Scientific Reports, 12, 22511.",
        "https://www.nature.com/articles/s41598-022-26413-6",
    ),
    (
        "Pajic, P., Landau, L., Gokcumen, O., & Ruhl, S. (2025). Saliva Protein Genes in Humans were Shaped during Primate Evolution. Genome Biology and Evolution, 17(9), evaf165.",
        "https://doi.org/10.1093/gbe/evaf165",
    ),
    (
        "Mase, M., Viziano, A., Strapazzon, G., Alessandrini, M., & Micarelli, A. (2023). Auditory function in humans at high altitude: A scoping review. PLOS ONE, 18(9), e0291919.",
        "https://pmc.ncbi.nlm.nih.gov/articles/PMC10513325/",
    ),
    (
        "Fishbein, A. B., Knutson, K. L., & Zee, P. C. (2021). Circadian disruption and human health. Journal of Clinical Investigation, 131(19), e148286.",
        "https://www.jci.org/articles/view/148286",
    ),
    (
        "Mishrikoti, P. V., Prarthana, H. M., & Lamani, T. (2024). Role of Dinacharya to maintain circadian rhythm for cell rejuvenation: A review. Journal of Ayurveda and Integrated Medical Sciences.",
        "https://jaims.in/index.php/jaims/article/view/3520",
    ),
    (
        "Costello, D. M., et al. (2022). A review of simulation training and new 3D computer-generated synthetic organs for robotic surgery education. Journal of Robotic Surgery, 16, 749-763.",
        "https://pure.johnshopkins.edu/en/publications/a-review-of-simulation-training-and-new-3d-computer-generated-syn/",
    ),
    (
        "Schwarz, P., Hellmers, S., Spanknebel, S., Hurlemann, R., & Hein, A. (2024). Humanoid patient robot for diagnostic training in medical and psychiatric education. Frontiers in Robotics and AI, 11, 1424845.",
        "https://www.frontiersin.org/journals/robotics-and-ai/articles/10.3389/frobt.2024.1424845/full",
    ),
    (
        "Murphy, D., Mitchell, J., Vacher, J., Morley, E. J., & Puzzo, I. (2021). Wechsler Adult Intelligence Scale full scale IQ of male admissions to a high secure psychiatric hospital over six decades. International Journal of Forensic Mental Health, 20(4), 386-397.",
        "https://paloaltou.edu/resources/translating-research-into-practice-blog/a-reverse-flynn-effect-trends-in-six-decades-of-neuropsychological-data-in-a-uk-high-security-population",
    ),
    (
        "Better Than Us. (2018-2019). Russian science fiction television series.",
        "https://en.wikipedia.org/wiki/Better_Than_Us",
    ),
    (
        "Connes, A. (1994). Noncommutative Geometry. Academic Press.",
        "https://alainconnes.org/wp-content/uploads/book94bigpdf.pdf",
    ),
    (
        "Smith, S. P. (2011). The space of Penrose tilings and the non-commutative curve with homogeneous coordinate ring k<x,y>/(y^2). arXiv:1104.3811.",
        "https://arxiv.org/abs/1104.3811",
    ),
    (
        "Nielsen, J. L. (forthcoming). The Topological Unified Field Theory on the Complex Hopf Fibration. International Journal of Topology. PhilPapers/PhilArchive record.",
        "https://philpapers.org/rec/NIETTU",
    ),
    (
        "Jaffe, A., & Liu, Z. (n.d.). A Mathematical Picture Language Program.",
        "https://arthurjaffe.com/Assets/pdf/PictureLanguage.pdf",
    ),
    (
        "Bourgade, P., & Huang, J. (2026). Loop Equations Characterize Random Matrix Statistics. arXiv:2607.07617.",
        "https://arxiv.org/abs/2607.07617",
    ),
    (
        "Axelrod, R., & Hamilton, W. D. (1981). The evolution of cooperation. Science, 211(4489), 1390-1396.",
        "https://pubmed.ncbi.nlm.nih.gov/7466396/",
    ),
    (
        "Wu, J., & Axelrod, R. (1995). How to cope with noise in the iterated prisoner's dilemma. Journal of Conflict Resolution, 39(1), 183-189.",
        "https://doi.org/10.1177/0022002795039001008",
    ),
    (
        "Veritasium. (2024, January 15). What The Prisoner's Dilemma Reveals About Life, The Universe, and Everything.",
        "https://www.veritasium.com/videos/2024/1/15/what-the-prisoners-dilemma-reveals-about-life-the-universe-and-everything",
    ),
    (
        "The Unicode Consortium. Unicode Bidirectional Algorithm, Unicode Standard Annex #9.",
        "https://www.unicode.org/reports/tr9/",
    ),
    (
        "World Wide Web Consortium. Character Model for the World Wide Web 1.0: Fundamentals.",
        "https://www.w3.org/TR/charmod/",
    ),
    (
        "Toosarvandani, M. (2021). Remembering Language. The Humanities Institute, UC Santa Cruz.",
        "https://thi.ucsc.edu/memory-series-maziar-toosarvandani/",
    ),
    (
        "Austrian Centre for Digital Humanities and Cultural Heritage. Romani Language Database.",
        "https://www.oeaw.ac.at/acdh/research/linguistics/resources/structured-datasets/romani-language-database",
    ),
    (
        "cppreference.com. Coroutines (C++ language reference).",
        "https://en.cppreference.com/w/cpp/language/coroutines",
    ),
    (
        "Aygp-dr. Liquid Neural Networks (LNN) implementation repository. GitHub. Draft/in-progress software artifact.",
        "https://github.com/aygp-dr/liquid-neural-networks",
    ),
    (
        "Leike, J., Martic, M., Krakovna, V., Ortega, P. A., Everitt, T., Lefrancq, A., Orseau, L., & Legg, S. (2017). AI Safety Gridworlds. arXiv:1711.09883.",
        "https://arxiv.org/abs/1711.09883",
    ),
    (
        "Grandjean, M. (vector), McGeddon (picture), & U.S. Air Force hit-plot concept. (2021). Survivorship-bias.svg. Wikimedia Commons. CC BY-SA 4.0.",
        "https://commons.wikimedia.org/wiki/File:Survivorship-bias.svg",
    ),
    (
        "Smith, J. E., Zuo, M., Kuhlke, W., Sprinkle, B., & Ristroph, L. (2026). Geometry controls momentum flux in the sprinkler problem. Proceedings of the National Academy of Sciences, 123(30).",
        "https://doi.org/10.1073/pnas.2537479123",
    ),
    (
        "Faculty of Science, The University of Hong Kong. (2026). Internationally Renowned Mathematician Professor Ha Van Vu Joins HKU.",
        "https://www.scifac.hku.hk/news/prof-van-h-vu-joins-hku",
    ),
    (
        "Satzer, B. (2020). Review of Qiang Du, Nonlocal Modeling, Analysis, and Computation. MAA Reviews.",
        "https://www.columbia.edu/~qd2125/MAA-review.pdf",
    ),
    (
        "Carlini, E., & Tozza, S. (2024). A scheme for the game p-Laplacian and its application to image inpainting. Applied Mathematics and Computation, 461, 128299.",
        "https://doi.org/10.1016/j.amc.2023.128299",
    ),
    (
        "Korner, T. W. (2018). Coding and Cryptography. University of Cambridge lecture notes.",
        "https://www.dpmms.cam.ac.uk/~twk10/Shan.pdf",
    ),
    (
        "Malone, J. L. (n.d.). Sequential Depletion Ordering with Residual-Fraction Costs. Local manuscript supplied for capstone stress testing.",
        "C:/Users/anaxe/Downloads/sequential_depletion_ordering.pdf",
    ),
    (
        "Ferguson, T. S. The Kelly Betting System for Favorable Games. UCLA Statistics lecture notes.",
        "https://www.math.ucla.edu/~tom/stat596/Kelly.pdf",
    ),
    (
        "Hughes, J., & Haran, M. (2010). Dimension Reduction and Alleviation of Confounding for Spatial Generalized Linear Mixed Models. arXiv:1011.6649.",
        "https://arxiv.org/abs/1011.6649",
    ),
    (
        "Loh, P.-R., Tucker, G., Bulik-Sullivan, B. K., et al. (2015). Efficient Bayesian mixed-model analysis increases association power in large cohorts. Nature Genetics, 47, 284-290.",
        "https://doi.org/10.1038/ng.3190",
    ),
    (
        "Loh, P.-R., Lipson, M., Patterson, N., et al. (2013). Inferring admixture histories of human populations using linkage disequilibrium. Genetics, 193(4), 1233-1254.",
        "https://arxiv.org/abs/1211.0251",
    ),
    (
        "Lipson, M., Loh, P.-R., Levin, A., Reich, D., Patterson, N., & Berger, B. (2013). Efficient moment-based inference of admixture parameters and sources of gene flow. Molecular Biology and Evolution, 30(8), 1788-1802.",
        "https://arxiv.org/abs/1212.2555",
    ),
    (
        "Ellenberg, J. (2014). Popular explanation of Maryam Mirzakhani's dynamics-and-geometry work, quoted in Maryam Mirzakhani biography.",
        "https://en.wikipedia.org/wiki/Maryam_Mirzakhani",
    ),
    (
        "International Mathematical Union. (2014). The Work of Maryam Mirzakhani.",
        "https://www.mathunion.org/fileadmin/IMU/Prizes/Fields/2014/news_release_mirzakhani.pdf",
    ),
    (
        "Wolpert, S. A. (2011). Mirzakhani's volume recursion and approach for the Witten-Kontsevich theorem on moduli tautological intersection numbers. arXiv:1108.0174.",
        "https://arxiv.org/abs/1108.0174",
    ),
    (
        "Bellare, M., Goldreich, O., & Petrank, E. Uniform Generation of NP-witnesses using an NP-oracle.",
        "https://cseweb.ucsd.edu/~mihir/papers/ug.pdf",
    ),
    (
        "Bailey, D. H., & Borwein, J. M. (2020). PSLQ: An Algorithm to Discover Integer Relations.",
        "https://www.davidhbailey.com/dhbpapers/pslq-comp-alg.pdf",
    ),
    (
        "Bailey, D. H., Borwein, J. M., & Girgensohn, R. (1994). Experimental evaluation of Euler sums. Experimental Mathematics, 3(1), 17-30.",
        "https://projecteuclid.org/journals/experimental-mathematics/volume-3/issue-1/Experimental-evaluation-of-Euler-sums/em/1062621000.full",
    ),
    (
        "Zhang, T. (2014, revised 2018). A note on the non-commutative arithmetic-geometric mean inequality. arXiv:1411.5058.",
        "https://arxiv.org/abs/1411.5058",
    ),
    (
        "Recht, B., & Re, C. (2012). Toward a Noncommutative Arithmetic-Geometric Mean Inequality: Conjectures, Case-studies, and Consequences. PMLR, 23, 11.1-11.24.",
        "https://proceedings.mlr.press/v23/recht12.html",
    ),
    (
        "Lai, Z., & Lim, L.-H. (2020). Recht-Re Noncommutative Arithmetic-Geometric Mean Conjecture is False. arXiv:2006.01510.",
        "https://arxiv.org/abs/2006.01510",
    ),
    (
        "Harvard Health Publishing. (2026). What is a normal heart rate?",
        "https://www.health.harvard.edu/heart-health/what-your-heart-rate-is-telling-you",
    ),
    (
        "Erin Brockovich. (2000). Film dramatization of the Hinkley groundwater contamination legal case.",
        "https://en.wikipedia.org/wiki/Erin_Brockovich_(film)",
    ),
    (
        "Dwork, C., & Roth, A. (2014). The Algorithmic Foundations of Differential Privacy. Foundations and Trends in Theoretical Computer Science, 9(3-4), 211-407.",
        "https://doi.org/10.1561/0400000042",
    ),
    (
        "Sheffield, S. 18.175 Lecture 10: Zero-one laws and maximal inequalities. MIT.",
        "https://math.mit.edu/~sheffield/175/Lecture10.pdf",
    ),
    (
        "Karlin, A. R., & Peres, Y. (2016). Game Theory, Alive. American Mathematical Society.",
        "https://homes.cs.washington.edu/~karlin/GameTheoryBook.pdf",
    ),
    (
        "Morreau, M. (2014). Arrow's Theorem. Stanford Encyclopedia of Philosophy.",
        "https://plato.stanford.edu/entries/arrows-theorem/",
    ),
    (
        "Arrow, K. J. (1950). A Difficulty in the Concept of Social Welfare. Journal of Political Economy, 58(4), 328-346.",
        "https://doi.org/10.1086/256963",
    ),
    (
        "Park, J.-W., Kim, J. U., Ghim, C.-M., & Kim, C. U. (2022). The Boltzmann fair division for distributive justice. Scientific Reports, 12, 16179.",
        "https://doi.org/10.1038/s41598-022-19792-3",
    ),
    (
        "Asad, R., Babanezhad, R., Laradji, I., Le Roux, N., & Vaswani, S. (2024). Fast Convergence of Softmax Policy Mirror Ascent for Bandits & Tabular MDPs. OPT 2024.",
        "https://opt-ml.org/papers/2024/paper83.pdf",
    ),
    (
        "Harchol-Balter, M. (2013). Performance Modeling and Design of Computer Systems: Queueing Theory in Action. Cambridge University Press.",
        "https://www.cs.cmu.edu/~harchol/PerformanceModeling/book.html",
    ),
    (
        "Blei, D. M., Ng, A. Y., & Jordan, M. I. (2003). Latent Dirichlet Allocation. Journal of Machine Learning Research, 3, 993-1022.",
        "https://www.jmlr.org/papers/v3/blei03a.html",
    ),
    (
        "Charles River Analytics. Figaro Programming Language and Core Libraries. GitHub.",
        "https://github.com/charles-river-analytics/figaro",
    ),
]


def run_stress_test() -> dict:
    env = os.environ.copy()
    env["PYTHONPATH"] = str(WORK / ".packages")
    result = subprocess.run(
        [str(PYTHON), str(PROJECT / "run_evidence_gate_stress_test.py")],
        cwd=str(ROOT),
        env=env,
        check=True,
        capture_output=True,
        text=True,
    )
    return json.loads(result.stdout)


def run_zdd_audit() -> dict:
    env = os.environ.copy()
    env["PYTHONPATH"] = str(PROJECT) + os.pathsep + str(WORK / ".packages")
    result = subprocess.run(
        [str(PYTHON), str(PROJECT / "zdd_sparse_claim_family.py")],
        cwd=str(ROOT),
        env=env,
        check=True,
        capture_output=True,
        text=True,
    )
    return json.loads(result.stdout)


def run_model_selection_audit() -> dict:
    env = os.environ.copy()
    env["PYTHONPATH"] = str(PROJECT) + os.pathsep + str(WORK / ".packages")
    result = subprocess.run(
        [str(PYTHON), str(PROJECT / "model_selection_claim_policy.py")],
        cwd=str(ROOT),
        env=env,
        check=True,
        capture_output=True,
        text=True,
    )
    return json.loads(result.stdout)


def run_depletion_ordering_experiment() -> dict:
    env = os.environ.copy()
    env["PYTHONPATH"] = str(PROJECT) + os.pathsep + str(WORK / ".packages")
    result = subprocess.run(
        [str(PYTHON), str(PROJECT / "depletion_ordering_experiment.py")],
        cwd=str(ROOT),
        env=env,
        check=True,
        capture_output=True,
        text=True,
    )
    return json.loads(result.stdout)


def run_orthogonal_projection_experiment() -> dict:
    env = os.environ.copy()
    env["PYTHONPATH"] = str(PROJECT) + os.pathsep + str(WORK / ".packages")
    result = subprocess.run(
        [str(PYTHON), str(PROJECT / "orthogonal_projection_experiment.py")],
        cwd=str(ROOT),
        env=env,
        check=True,
        capture_output=True,
        text=True,
    )
    return json.loads(result.stdout)


def run_kelly_runtime_budget_experiment() -> dict:
    env = os.environ.copy()
    env["PYTHONPATH"] = str(PROJECT) + os.pathsep + str(WORK / ".packages")
    result = subprocess.run(
        [str(PYTHON), str(PROJECT / "kelly_runtime_budget_experiment.py")],
        cwd=str(ROOT),
        env=env,
        check=True,
        capture_output=True,
        text=True,
    )
    return json.loads(result.stdout)


def run_uniform_witness_sampling_experiment() -> dict:
    env = os.environ.copy()
    env["PYTHONPATH"] = str(PROJECT) + os.pathsep + str(WORK / ".packages")
    result = subprocess.run(
        [str(PYTHON), str(PROJECT / "uniform_witness_sampling_experiment.py")],
        cwd=str(ROOT),
        env=env,
        check=True,
        capture_output=True,
        text=True,
    )
    return json.loads(result.stdout)


def run_consent_aggregation_experiment() -> dict:
    env = os.environ.copy()
    env["PYTHONPATH"] = str(PROJECT) + os.pathsep + str(WORK / ".packages")
    result = subprocess.run(
        [str(PYTHON), str(PROJECT / "consent_aggregation_experiment.py")],
        cwd=str(ROOT),
        env=env,
        check=True,
        capture_output=True,
        text=True,
    )
    return json.loads(result.stdout)


def run_event_algebra_experiment() -> dict:
    env = os.environ.copy()
    env["PYTHONPATH"] = str(PROJECT) + os.pathsep + str(WORK / ".packages")
    result = subprocess.run(
        [str(PYTHON), str(PROJECT / "event_algebra_experiment.py")],
        cwd=str(ROOT),
        env=env,
        check=True,
        capture_output=True,
        text=True,
    )
    return json.loads(result.stdout)


def run_measure_on_measures_experiment() -> dict:
    env = os.environ.copy()
    env["PYTHONPATH"] = str(PROJECT) + os.pathsep + str(WORK / ".packages")
    result = subprocess.run(
        [str(PYTHON), str(PROJECT / "measure_on_measures_experiment.py")],
        cwd=str(ROOT),
        env=env,
        check=True,
        capture_output=True,
        text=True,
    )
    return json.loads(result.stdout)


def run_mahalanobis_covariate_experiment() -> dict:
    env = os.environ.copy()
    env["PYTHONPATH"] = str(PROJECT) + os.pathsep + str(WORK / ".packages")
    result = subprocess.run(
        [str(PYTHON), str(PROJECT / "mahalanobis_covariate_experiment.py")],
        cwd=str(ROOT),
        env=env,
        check=True,
        capture_output=True,
        text=True,
    )
    return json.loads(result.stdout)


def run_noncommutative_amgm_counterexample() -> dict:
    env = os.environ.copy()
    env["PYTHONPATH"] = str(PROJECT) + os.pathsep + str(WORK / ".packages")
    result = subprocess.run(
        [str(PYTHON), str(PROJECT / "noncommutative_amgm_counterexample.py")],
        cwd=str(ROOT),
        env=env,
        check=True,
        capture_output=True,
        text=True,
    )
    return json.loads(result.stdout)


def run_tail_maximal_inequality_experiment() -> dict:
    env = os.environ.copy()
    env["PYTHONPATH"] = str(PROJECT) + os.pathsep + str(WORK / ".packages")
    result = subprocess.run(
        [str(PYTHON), str(PROJECT / "tail_maximal_inequality_experiment.py")],
        cwd=str(ROOT),
        env=env,
        check=True,
        capture_output=True,
        text=True,
    )
    return json.loads(result.stdout)


def run_hex_boundary_invariant_experiment() -> dict:
    env = os.environ.copy()
    env["PYTHONPATH"] = str(PROJECT) + os.pathsep + str(WORK / ".packages")
    result = subprocess.run(
        [str(PYTHON), str(PROJECT / "hex_boundary_invariant_experiment.py")],
        cwd=str(ROOT),
        env=env,
        check=True,
        capture_output=True,
        text=True,
    )
    return json.loads(result.stdout)


def run_hex_scaling_coarse_grain_experiment() -> dict:
    env = os.environ.copy()
    env["PYTHONPATH"] = str(PROJECT) + os.pathsep + str(WORK / ".packages")
    result = subprocess.run(
        [str(PYTHON), str(PROJECT / "hex_scaling_coarse_grain_experiment.py")],
        cwd=str(ROOT),
        env=env,
        check=True,
        capture_output=True,
        text=True,
    )
    return json.loads(result.stdout)


def run_proof_status_poset_experiment() -> dict:
    env = os.environ.copy()
    env["PYTHONPATH"] = str(PROJECT) + os.pathsep + str(WORK / ".packages")
    result = subprocess.run(
        [str(PYTHON), str(PROJECT / "proof_status_poset_experiment.py")],
        cwd=str(ROOT),
        env=env,
        check=True,
        capture_output=True,
        text=True,
    )
    return json.loads(result.stdout)


def run_picture_language_diagram_audit() -> dict:
    env = os.environ.copy()
    env["PYTHONPATH"] = str(PROJECT) + os.pathsep + str(WORK / ".packages")
    result = subprocess.run(
        [str(PYTHON), str(PROJECT / "picture_language_diagram_audit.py")],
        cwd=str(ROOT),
        env=env,
        check=True,
        capture_output=True,
        text=True,
    )
    return json.loads(result.stdout)


def run_loop_equation_runtime_stability() -> dict:
    env = os.environ.copy()
    env["PYTHONPATH"] = str(PROJECT) + os.pathsep + str(WORK / ".packages")
    result = subprocess.run(
        [str(PYTHON), str(PROJECT / "loop_equation_runtime_stability.py")],
        cwd=str(ROOT),
        env=env,
        check=True,
        capture_output=True,
        text=True,
    )
    return json.loads(result.stdout)


def run_branch_factor_path_stability() -> dict:
    env = os.environ.copy()
    env["PYTHONPATH"] = str(PROJECT) + os.pathsep + str(WORK / ".packages")
    result = subprocess.run(
        [str(PYTHON), str(PROJECT / "branch_factor_path_stability.py")],
        cwd=str(ROOT),
        env=env,
        check=True,
        capture_output=True,
        text=True,
    )
    return json.loads(result.stdout)


def export_pdf(docx_path: Path) -> int:
    pdf_path = docx_path.with_suffix(".pdf")
    result = subprocess.run(
        [str(PYTHON), str(WORK / "export_docx_with_word.py"), str(docx_path), str(pdf_path)],
        cwd=str(ROOT),
        check=True,
        capture_output=True,
        text=True,
    )
    for line in result.stdout.splitlines():
        if line.startswith("WordPages="):
            return int(line.split("=", 1)[1])
    return 0


def sha256(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as f:
        for chunk in iter(lambda: f.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_grid = tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        tbl.append(tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        tbl_grid.append(col)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            set_cell_width(cell, widths[i])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            margins = tc_pr.find(qn("w:tcMar"))
            if margins is None:
                margins = OxmlElement("w:tcMar")
                tc_pr.append(margins)
            for side, value in {"top": "80", "bottom": "80", "start": "120", "end": "120"}.items():
                node = margins.find(qn(f"w:{side}"))
                if node is None:
                    node = OxmlElement(f"w:{side}")
                    margins.append(node)
                node.set(qn("w:w"), value)
                node.set(qn("w:type"), "dxa")


def paragraph_border_bottom(paragraph, color="D7DBE2") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def style_doc(doc: Document, title: str, subtitle: str, doc_type: str) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10.8)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.22

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 11, 5),
        ("Heading 3", 12, DARK, 8, 3),
    ]:
        st = styles[style_name]
        st.font.name = "Calibri"
        st._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        st._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        st.font.size = Pt(size)
        st.font.color.rgb = color
        st.font.bold = True
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)

    header_p = section.header.paragraphs[0]
    header_p.text = f"{doc_type} | Evidence-Gated Agentic Medical LLMs"
    header_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in header_p.runs:
        run.font.size = Pt(8.5)
        run.font.color.rgb = MUTED
    footer_p = section.footer.paragraphs[0]
    footer_p.text = "Ravi Bajaj | Module 14 Final Capstone"
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer_p.runs:
        run.font.size = Pt(8.5)
        run.font.color.rgb = MUTED

    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_after = Pt(1)
    r = kicker.add_run(doc_type.upper())
    r.bold = True
    r.font.size = Pt(8.5)
    r.font.color.rgb = BLUE

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = INK

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(subtitle)
    r.font.size = Pt(11.5)
    r.font.color.rgb = MUTED

    meta = doc.add_paragraph()
    meta.paragraph_format.space_after = Pt(8)
    r = meta.add_run("Author: ")
    r.bold = True
    meta.add_run("Ravi Bajaj | Target venue: ML4H 2026 Proceedings Track | Course: AI in Healthcare")

    rule = doc.add_paragraph()
    paragraph_border_bottom(rule)


def add_para(doc: Document, text: str, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text[len(bold_prefix) :])
    else:
        p.add_run(text)


def add_code_block(doc: Document, code: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.right_indent = Inches(0.25)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "EEF3F8")
    p_pr.append(shd)
    for i, line in enumerate(code.splitlines()):
        if i:
            p.add_run().add_break()
        r = p.add_run(line)
        r.font.name = "Consolas"
        r._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
        r._element.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")
        r.font.size = Pt(8.8)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        p.add_run(item)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(3)
        p.add_run(item)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, LIGHT)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(9)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(value)
            r.font.size = Pt(8.8)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_result_matrix(doc: Document, rows: list[dict[str, str]]) -> None:
    compact = []
    for row in rows:
        compact.append(
            [
                row["scenario_id"],
                row["label"],
                row["endpoint_type"].replace("_", " "),
                ACTION_SHORT.get(row["actual_action"], row["actual_action"].replace("_", " ").lower()),
            ]
        )
    midpoint = (len(compact) + 1) // 2
    add_table(doc, ["ID", "Scenario", "Endpoint", "Runtime action"], compact[:midpoint], [650, 4100, 2100, 2510])
    add_table(doc, ["ID", "Scenario", "Endpoint", "Runtime action"], compact[midpoint:], [650, 4100, 2100, 2510])


def pair_rows(rows: list[list[str]]) -> list[list[str]]:
    midpoint = (len(rows) + 1) // 2
    left = rows[:midpoint]
    right = rows[midpoint:]
    paired = []
    for index, left_row in enumerate(left):
        right_row = right[index] if index < len(right) else ["", ""]
        paired.append([left_row[0], left_row[1], right_row[0], right_row[1]])
    return paired


def add_picture(doc: Document, path: Path, width: float, caption: str) -> None:
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(8)
    r = cap.add_run(caption)
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = MUTED


def source_rows(limit: int | None = None) -> list[list[str]]:
    srcs = SOURCES if limit is None else SOURCES[:limit]
    return [[str(i + 1), citation, url] for i, (citation, url) in enumerate(srcs)]


def sources_matching(needles: list[str]) -> list[tuple[str, str]]:
    selected: list[tuple[str, str]] = []
    for needle in needles:
        match = next(
            ((citation, url) for citation, url in SOURCES if needle.lower() in (citation + " " + url).lower()),
            None,
        )
        if match and match not in selected:
            selected.append(match)
    return selected


def create_proposal(
    summary: dict,
    zdd_summary: dict,
    model_summary: dict,
    depletion_summary: dict,
    projection_summary: dict,
    kelly_summary: dict,
    witness_summary: dict,
    consent_summary: dict,
    event_summary: dict,
    measure_summary: dict,
    mahalanobis_summary: dict,
    hex_scaling_summary: dict,
    proof_status_summary: dict,
    picture_summary: dict,
    loop_summary: dict,
    branch_summary: dict,
) -> Path:
    path = PACKAGE / "proposal" / "Module_14_Capstone_Proposal_Evidence_Gated_Medical_LLMs_Ravi_Bajaj.docx"
    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    style_doc(
        doc,
        "Surrogate-Aware Runtime Assurance for Agentic Medical LLMs",
        "Preventing evidence overclaiming in clinician-facing AI systems",
        "Research Proposal",
    )

    doc.add_heading("Problem and Hypothesis", level=1)
    add_para(
        doc,
        "Medical LLMs are often tested for hallucination, but agentic workflows create a subtler failure: a real citation, workflow metric, surrogate endpoint, population-mismatched result, or privacy-limited genetic release can be converted into a stronger patient-outcome claim than the evidence supports. The practical target is recurring hospital genetics: pharmacogenomic alerts, newborn-screening follow-up, ACMG secondary findings, hereditary cancer/FH/Lynch flags, VUS handling, and family-linked privacy boundaries. Hypothesis: a runtime assurance monitor can reduce unsafe clinical overclaiming by gating each LLM output or proposed action by citation validity, endpoint strength, study design, population fit, context, opportunity cost, clinician authority, privacy budget, and consent boundary.",
    )

    doc.add_heading("Key Background", level=1)
    add_bullets(
        doc,
        [
            "Stegenga, Medical Nihilism: motivates caution when evidentiary pipelines overstate treatment benefit.",
            "Byrne, Domenico, and Moore: pragmatic patient-level RCTs remain the standard for AI patient-outcome claims.",
            "FDA-NIH BEST / FDA surrogate endpoint materials: distinguish clinical outcomes, validated surrogates, ordinary surrogates, and biomarkers.",
            "CPIC, ACMG, HRSA/RUSP, ACT sheets, and CDC Tier 1 materials anchor the project in hospital-facing genetics rather than speculative genetics.",
            "Dwork and Roth: differential privacy supplies measurable privacy loss, composition, post-processing, and group/family privacy concepts.",
            "Loh-style statistical genetics, Arrow social choice, Harchol-Balter scheduling, LDA, and Mahalanobis distance motivate population-structure, consent-aggregation, queueing, mixture-model, and covariate-transport gates.",
            "Hou, Ji, Zhang, and Stefanidis on Kolmogorov-Arnold Networks: architecture/theorem claims can guide model selection, but remain method evidence rather than clinical permission.",
        ],
    )

    doc.add_heading("Method and Stage 1 Results", level=1)
    add_para(
        doc,
        f"Stage 1 is implemented as a deterministic synthetic assurance harness with {summary['scenario_count']} scenarios and {summary['matched_expected_count']} matched runtime decisions. The controller may allow bounded second-reader use, narrow to surrogate/process language, allow simulation-only rehearsal, require provenance/proof-status or consent audit, stress-test generalization/order/confounding, or deny the claim. A ZDD audit compactly represents admissible sparse evidence states; a model-selection layer chooses among {model_summary['candidate_model_count']} claim models; all code and JSON/CSV results are included.",
    )
    add_bullets(
        doc,
        [
            f"ZDD/model selection/proof status: {zdd_summary['observed_family_size']} observed subsets over {zdd_summary['feature_universe_size']} features; {model_summary['matched_expected_count']}/{model_summary['scenario_count']} actions matched; the KAN architecture case was narrowed to method evidence; the proof-status poset allowed hard-outcome permission in {proof_status_summary['permission_counts']['hard_outcome_allowed_with_caveats']}/{proof_status_summary['states_enumerated']} states and kept NIETTU audit-only.",
            f"Runtime resources: {depletion_summary['permutation_count']} orderings tested; heterogeneous resources showed context reversal. Kelly-style exposure capped claim authority after all-in ruin probability {kelly_summary['all_in_ruin_probability']:.3f}.",
            f"Confounding and witnesses: orthogonal projection reduced max design inner product to {projection_summary['max_abs_design_inner_product_after']:.2e}; biased strongest-five witness selection had variation distance {witness_summary['variation_distance_from_uniform']:.3f}.",
            f"Hospital genetics: the harness includes recurring genetics-result triage; family preferences produced a cycle, full raw release was blocked, and partial inclusion stayed runtime-safe with composed epsilon {consent_summary['partial_plus_audit_composed_loss']['epsilon']:.2f}.",
            f"Population/source/boundary transport: event algebra verified measurable identities; measure-on-measures max TV was {measure_summary['max_source_tv_to_mixture']:.3f}; Mahalanobis flagged {', '.join(mahalanobis_summary['flagged_sources'])}; Hex scaling sampled {hex_scaling_summary['sampled_full_boards']} boards with {hex_scaling_summary['ambiguous_terminal_count']} ambiguous terminals and found generic smoothing could flip the boundary at rate {hex_scaling_summary['max_coarse_grain_flip_rate']:.3f}.",
            f"Multimodal and universality transport: picture-language audit classified {picture_summary['diagram_count']} diagrammatic artifacts with zero hard-outcome permissions; loop-equation stability kept the stable transfer chain at error {loop_summary['stable_final_error']:.4f} under budget {loop_summary['stable_budget']:.3f}; switching-cumulant checks canceled the main quadratic term and kept rare-event/replacement errors inside budget; branch-factor audit allowed the separated synthetic branch but blocked the near-collision family/population branch.",
        ],
    )

    doc.add_heading("Value and Stage 2", level=1)
    add_para(
        doc,
        "The proposal contributes an auditable standard for permissible LLM inputs and outputs: fluent text cannot upgrade evidence, and family-linked genetic features cannot be treated as ordinary single-person records. Stage 2 would apply the abstract gates to faculty-approved hospital genetics workflows first, then only later to broader polygenic-risk, wearable/telemedicine, or multilingual documentation settings. The monitor preserves clinician authority and requires patient-outcome evidence before outcome claims.",
    )

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("Selected References")
    r.bold = True
    r.font.size = Pt(8.5)
    r.font.color.rgb = DARK
    proposal_sources = [
        (
            "Stegenga (2018), Medical Nihilism.",
            "https://global.oup.com/academic/product/medical-nihilism-9780198747048",
        ),
        (
            "Byrne, Domenico, & Moore (2024), pragmatic RCTs for AI outcomes.",
            "https://kjronline.org/DOIx.php?id=10.3348%2Fkjr.2023.1016",
        ),
        (
            "FDA-NIH BEST Resource (2016-2025), biomarkers and endpoints.",
            "https://www.ncbi.nlm.nih.gov/books/NBK326791/",
        ),
        (
            "Dwork & Roth (2014), Algorithmic Foundations of Differential Privacy.",
            "https://doi.org/10.1561/0400000042",
        ),
        (
            "Bourgade & Huang (2026), Loop Equations Characterize Random Matrix Statistics.",
            "https://arxiv.org/abs/2607.07617",
        ),
    ]
    for citation, url in proposal_sources:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(citation + " ")
        r.font.size = Pt(7.2)
        u = p.add_run(url)
        u.font.size = Pt(7.2)
        u.font.color.rgb = BLUE

    doc.save(path)
    return path


def create_final_paper(
    summary: dict,
    zdd_summary: dict,
    model_summary: dict,
    depletion_summary: dict,
    projection_summary: dict,
    kelly_summary: dict,
    witness_summary: dict,
    consent_summary: dict,
    event_summary: dict,
    measure_summary: dict,
    mahalanobis_summary: dict,
    matrix_summary: dict,
    tail_summary: dict,
    hex_summary: dict,
    hex_scaling_summary: dict,
    proof_status_summary: dict,
    picture_summary: dict,
    loop_summary: dict,
    branch_summary: dict,
) -> Path:
    path = PACKAGE / "paper" / "Module_14_Final_Paper_Evidence_Gated_Medical_LLMs_Ravi_Bajaj.docx"
    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    style_doc(
        doc,
        "Surrogate-Aware Runtime Assurance for Agentic Medical LLMs",
        "A synthetic stress-test framework for evidence-gated clinical claims",
        "Final Paper",
    )

    doc.add_heading("Abstract", level=1)
    add_para(
        doc,
        "Medical large language models (LLMs) are often evaluated for hallucination and citation accuracy, but agentic clinical workflows introduce a subtler failure: evidence overclaiming by systems that retrieve, summarize, plan, and draft clinical-sounding actions. Building on prior course evaluations of LLMs, decision support, automation, and epidemic simulation, this paper proposes a runtime-assurance framework that gates LLM outputs and proposed actions by citation validity, endpoint strength, study design, target-population fit, actionability, opportunity cost, clinician authority, source provenance, runtime budget order, confounding, proportional claim exposure, witness-selection bias, privacy loss, and consent aggregation. A synthetic stress-test set distinguishes hard outcomes from surrogate/process endpoints, flags unverifiable citations, treats low clinician uptake as possible safety filtering, and requires abstention when context is missing. Additional probes cover multimodal telemedicine, consumer heart-rate/wearable physiology, cross-script and Romani-style language provenance, KAN-style architecture claims, SIK3/tinnitus genetics, saliva evolutionary genomics, altitude/circadian transport, robotic-cadaver rehearsal, neuropsychological norm drift, sequential depletion of scarce runtime resources, and family genomic consent under differential-privacy-style budget accounting. A model-selection audit then chooses the least unsafe claim form compatible with the available evidence. The framework is motivated by Medical Nihilism, AI Safety Gridworlds, Po-Ru Loh-style bioinformatics, differential privacy, social-choice limits, experimental mathematics, residual-depletion ordering, and pragmatic randomized-trial standards. It does not validate a clinical product, but provides a testable architecture for safer clinician-facing agentic medical LLMs.",
    )

    doc.add_heading("1. Introduction", level=1)
    add_para(
        doc,
        "The central problem for medical LLMs is not simply whether they can produce plausible clinical text. It is whether agentic workflows that retrieve, summarize, plan, rank, and draft actions can keep the evidentiary strength of a claim separate from the fluency of the answer. A model may cite a real study and still make an unsafe inference if it treats screening completion, detection rate, evolutionary-genomics mechanism, or workflow efficiency as proof of improved mortality, morbidity, function, or quality of life.",
    )
    add_para(
        doc,
        "This paper frames that problem as runtime assurance for bounded medical agency. The LLM is allowed to act as a bounded assistant, second reader, summarizer, retriever, or question generator only after a monitor has classified the evidence and set a maximum permission level for the response or proposed action. In this framing, abstention is not a defect in the interface. Abstention is the behavior that prevents local evidence from being inflated into global clinical authority.",
    )
    add_para(
        doc,
        "The Wald survivorship-bias aircraft diagram is a useful warning for medical LLM evaluation. If reviewers only inspect fluent responses that return with plausible citations, they may miss the more important denominator: prompts that should have abstained, evidence chains that never terminate in a patient outcome, populations excluded from validation, or contexts that disappeared during summarization. The missing marks are often where the safety system should reinforce itself.",
    )

    doc.add_heading("2. Background", level=1)
    doc.add_heading("2.1 Medical Nihilism and Endpoint Discipline", level=2)
    add_para(
        doc,
        "Stegenga's Medical Nihilism argues for caution in medicine because biases, malleable study designs, and selective evidentiary practices can overstate benefits and understate harms. The point is not that medicine cannot work. The point is that the path from evidence to treatment confidence is fragile. Medical LLMs make this fragility operationally important because they can compress a messy literature into a confident answer that sounds more settled than the evidence itself.",
    )
    add_para(
        doc,
        "FDA and FDA-NIH BEST materials provide the practical distinction needed for the runtime layer: clinical outcomes measure how people feel, function, or survive; surrogate endpoints can be useful but require validation and, in many cases, confirmatory evidence. This distinction is the backbone of the gate.",
    )

    doc.add_heading("2.2 Pragmatic RCTs and Workflow Reality", level=2)
    add_para(
        doc,
        "Byrne, Domenico, and Moore argue that patient-outcome claims about healthcare AI should be tested in real workflows through pragmatic randomized trials. The CLOT trial illustrates why this matters. It embedded a validated prognostic model into clinical workflow, yet HA-VTE rates were not reduced and recommendations were accepted only 25.8 percent of the time. A shallow interpretation calls that last-mile failure. A safer interpretation asks whether clinicians were filtering recommendations for reasons the model did not encode.",
    )
    add_para(
        doc,
        "This capstone therefore treats low acceptance, opportunity cost, and clinician concern as signals to be modeled rather than nuisances to be optimized away. The same logic applies to models such as ANA triage systems: accuracy can be valuable, but if the model diverts scarce specialist attention or hides uncertainty, it may not improve care.",
    )

    doc.add_heading("2.3 Claim Grammar and Bounded Ambiguity", level=2)
    add_para(
        doc,
        "The formal-language analogy is deliberately bounded. Gold-style learnability asks what can be identified from a stream of positive examples. Carroll presents formal language theory and Gold learnability as tools for evaluating grammar systems. Kanazawa's work on k-valued categorial grammars gives the useful parameter: a symbol may receive at most k categories. In this paper, a citation, observation, or trial result is an evidence asset that may receive only a bounded set of evidence categories. The runtime monitor must resolve or preserve that ambiguity before an LLM composes a clinical claim.",
    )
    add_para(
        doc,
        "Vishnikin and Okhotin's unique-category assignment result sharpens the intuition. Removing lexical ambiguity before composition can still leave expressive power. Analogously, a medical LLM does not need to erase nuance; it needs to prevent a single paper from being spent simultaneously as workflow evidence, surrogate evidence, and hard-outcome evidence.",
    )
    add_para(
        doc,
        "The word elasticity is used here in Kanazawa's inductive-inference sense, not in the rubber-band sense of knot theory or Morse-theoretic visualization. A separate probability analogy comes from concentration and superconcentration: a runtime monitor should notice when variability, ambiguity, or endpoint roughness is being hidden by a smooth summary.",
    )
    add_para(
        doc,
        "Van Vu's work in random graph theory and random matrix theory is included only as background for this stability intuition: high-dimensional AI systems often look smooth in aggregate while depending on spectral, concentration, and perturbation behavior that must be monitored rather than assumed away.",
    )
    add_para(
        doc,
        "The bioinformatics anchor is closer to Po-Ru Loh-style statistical genetics than to an ordinary chatbot benchmark. Mixed models, linkage disequilibrium, admixture inference, and population structure all teach the same lesson: hidden structure can make a superficially portable signal confounded, miscalibrated, or population-specific. A medical LLM that reads genetic or polygenic-risk evidence should therefore expose ancestry/population fit, imputation status, uncertainty, and latent-factor handling before converting the evidence into a medical claim.",
    )
    add_para(
        doc,
        "The clinical genetics scope is deliberately ordinary and hospital-facing. Hospitals repeatedly encounter medication-gene interactions such as CYP2C19 and clopidogrel, positive newborn-screening follow-up that points clinicians to ACT sheets and algorithms, reportable ACMG secondary findings, CDC Tier 1 applications such as hereditary breast and ovarian cancer, Lynch syndrome, and familial hypercholesterolemia, and variants of uncertain significance that must not be overinterpreted. These are monthly or yearly governance and decision-support problems in real systems. They are a better first test bed than speculative designer-baby, open-ended wellness, or unconstrained polygenic-risk claims.",
    )
    add_para(
        doc,
        "Dwork and Roth provide the formal privacy backbone. Differential privacy is a definition rather than one algorithm; privacy loss is quantified by parameters such as epsilon and delta; composition consumes privacy budget; group privacy matters for families; and post-processing does not increase privacy loss if the underlying release was already private. For medical genetics, this means the access-control question is not simply who sees the full record. The runtime monitor must track which family-linked genomic coordinates, phenotype summaries, or polygenic-risk features are released, at what granularity, under what cumulative privacy budget.",
    )
    add_para(
        doc,
        "Arrow's theorem adds a consent-aggregation warning. Ranked preferences from relatives over options such as full raw reuse, trait-specific partial inclusion, audit-only use, or no reuse cannot always be collapsed into one fair family ordering. In the proposed framework, ranked choice is useful input, but the runtime decision preserves individual vetoes, partial permissions, and audit logs. This is the graded-poset point: moving downward from a full family genomic release can happen along multiple incomparable paths, not along one universal chain.",
    )
    add_para(
        doc,
        "Boltzmann fair division and softmax policy mirror ascent are used only inside that boundary. Entropy-controlled or softmax-style allocation can prioritize limited audit attention, model-routing probability, or privacy budget among feasible options, and SPMA supplies a modern mirror-ascent reference for fast softmax policy updates in bandit and tabular MDP settings. Neither method overrides consent or converts a governance decision into clinical evidence.",
    )
    add_para(
        doc,
        "Uniform generation of NP-witnesses supplies a final evaluation analogy. An evidence witness is a concrete support pattern for a claim: endpoint, design, population, and provenance. A medical LLM evaluator that samples only the strongest-looking witnesses can overstate the system's reliability. The proposed audit therefore records witness-selection bias rather than treating cherry-picked examples as representative.",
    )
    add_para(
        doc,
        "The measure-theory language is used in a finite, audit-friendly sense. An event is a measurable set of runtime histories, for example the set where a surrogate claim is promoted to a hard-outcome claim, privacy budget is exceeded, or a consent boundary is collapsed. A population, hospital, family, or evidence source can then be represented as a probability measure over those histories. A measure on the space of measures represents uncertainty over which source distribution is active. This is the abstract form of subgroup and data-source robustness: the mixture may look safe while one source measure remains high-risk.",
    )
    add_para(
        doc,
        "Sheffield's lecture on zero-one laws and maximal inequalities supplies a compact probability lens for runtime assurance. A tail event is a long-run property whose membership does not change when finitely many early observations are altered. That is a warning against certifying an agentic medical LLM from a short clean prefix. Kolmogorov's maximal inequality adds a finite-horizon monitor: even when increments are mean-zero, the relevant safety question may be whether the cumulative process crossed a boundary at any time, not merely where it ended.",
    )
    add_para(
        doc,
        "Karlin and Peres's Hex and Y exposition supplies a discrete topological analogy for terminal safety states and safe coarse-graining. Hex is finite and progressively bounded, and a completed standard board has exactly one winning crossing. Their Y-board reduction also uses local majority recoloring of small triples while preserving the global crossing/Y property. The capstone uses that idea as a design invariant: after required evidence features are known, a runtime state should not be both permitted and blocked, or neither classified; and any privacy or genetics coarse-graining should preserve the boundary property it claims to support. The strategy-stealing lesson is also useful for agentic systems: an unlogged extra move or evidence feature can advantage the acting agent, so hidden moves must be treated as safety-relevant.",
    )
    add_para(
        doc,
        "Mahalanobis distance makes this operational for covariate balance. Instead of asking whether a source population seems similar by one variable at a time, the monitor measures covariance-adjusted distance from a reference vector of ancestry principal components, age, rural access, missingness, and privacy constraints. Large distance is a runtime transport warning: the source needs source-specific validation before its evidence is used as if it belonged to the reference distribution.",
    )
    add_para(
        doc,
        "Latent Dirichlet Allocation adds a familiar machine-learning instance of this viewpoint. LDA represents documents as random mixtures over latent topics. In the capstone, an evidence artifact can analogously be represented as a mixture over latent claim topics such as hard outcome, surrogate, workflow, privacy, or consent. The runtime monitor can use such mixtures for routing and triage, but the mixture itself is not a clinical endpoint.",
    )
    add_para(
        doc,
        "Figaro is a natural future substrate for this layer because it is a probabilistic programming language for expressing rich probabilistic models and applying reasoning algorithms to evidence. This submission keeps the executable artifact in dependency-light Python, but the same finite event algebra, measure-on-measures source model, and runtime-history gates could be recast in Figaro for richer probabilistic replay and audit.",
    )
    add_para(
        doc,
        "Connes's use of Penrose tilings in noncommutative geometry supplies a second bounded analogy. The classical quotient of all Penrose tilings by translation is pathological because local patches recur everywhere while global tilings remain nonidentical; the remedy is to keep relational structure rather than replacing the space by a trivial quotient. Medical evidence has a similar failure mode for LLMs: the same finite patch of words, citations, biomarkers, or workflow metrics may recur across settings, but quotienting away population, time, instrument version, clinical role, and action context can erase the structure that makes a claim safe or unsafe.",
    )
    add_para(
        doc,
        "Cohen's forcing story gives a third bounded logic analogy. In set theory, a statement may be independent of a chosen axiom system. In medical LLM evaluation, a patient-outcome claim may be independent of the evidence package currently available to the agent. The safe runtime response is to expose the missing assumption, request additional evidence, or abstain, rather than letting fluent language act as a new axiom.",
    )
    add_para(
        doc,
        "The nonlocal, inpainting, and coding-theory references sharpen the same point from three directions. Du's nonlocal modeling emphasizes that some systems depend on interactions at a distance or across memory; family genetics, longitudinal records, and social-language context have that character. Carlini and Tozza's game p-Laplacian inpainting work is a useful caution for clinical AI: reconstructing missing parts can be principled, but reconstructed content must remain labeled as reconstruction rather than observed evidence. Korner's coding notes add the decodability requirement: after compression, captioning, blurring, encryption, or summarization, the downstream system must still know which evidence object it is reading.",
    )
    add_para(
        doc,
        "Jaffe and Liu's mathematical picture-language program makes the multimodal version precise. They distinguish an abstract picture language L, a target reality R, and a simulation map S connecting the two. In medical AI, a screenshot, wound photo, imaging slice, waveform, captioned video, or diagram can be a powerful object in L. The runtime gate asks whether the simulation map to the clinical reality is explicit, consented, provenance-preserved, and clinically validated before the LLM treats the picture as evidence.",
    )
    add_para(
        doc,
        "Experimental mathematics supplies a complementary proof-status analogy. PSLQ and Euler-sum searches can identify striking candidate relations, but a discovered relation is not the same object as a proof, and an uncertain memory of a source is not the same object as provenance. Medical LLM literature mining has the same risk: a plausible relation, even one discovered by a powerful search procedure, must remain labeled as hypothesis or surrogate until proof, validation, or patient-outcome evidence exists.",
    )
    add_para(
        doc,
        "The noncommutative AM-GM counterexample probe makes this point more concrete. Recht and Re's optimization motivation, Zhang's four-matrix inequality, and the later Lai-Lim disproof of the broader Recht-Re conjecture sit in a literature where simulations can support claims that later fail under exact adversarial construction. In this package I include only a bounded exact-arithmetic check: a 2 by 2 rank-one projector example violates the candidate bound by the factor "
        + str(matrix_summary["violation_factor_exact"])
        + ". The clinical analogy is not about matrices; it is that medical LLM evaluation should search for small, exact failure modes instead of trusting smooth aggregate simulations.",
    )
    add_para(
        doc,
        "Sequential depletion ordering adds a runtime-budget analogy. If each safety action consumes a fraction of remaining clinician attention, audit time, compute, or provenance bandwidth, then the order of actions is not neutral. In a single-resource synthetic experiment, small-first order minimizes raw residual-fraction cost, while large-first order minimizes convex log-depletion risk. In heterogeneous multi-resource settings, the locally preferred order can reverse with the residual state. This is why the capstone treats FIFO as an audit-replay policy, not as a universal safety policy.",
    )
    add_para(
        doc,
        "Harchol-Balter's queueing-theory and scheduling work is the systems version of this point. A hospital-facing agentic LLM is a queueing system under load: it schedules documentation, provenance checks, clinician interruptions, model calls, and escalation paths. Heavy-tailed workloads and resource-allocation policies can change response time and safety behavior, so runtime assurance must monitor scheduling policy rather than assuming that all checks can happen eventually with the same effect.",
    )
    add_para(
        doc,
        "Consumer heart-rate guidance supplies the simplest wearable example. Harvard Health notes that resting heart rate varies across people and can be affected by stress, anxiety, hormones, medication, physical activity level, age, and fitness. A wearable stream is therefore useful monitoring evidence, but not a standalone diagnostic or exercise-prescription authority. The runtime monitor should preserve that distinction before a medical LLM converts a number into reassurance, alarm, or advice.",
    )

    doc.add_heading("3. Methods", level=1)
    add_para(
        doc,
        f"I implemented a deterministic synthetic stress-test harness with {summary['scenario_count']} scenarios derived from course assignments, instructor feedback, and the cited literature. Each scenario specifies the requested claim, endpoint type, study design, citation status, population fit, context, opportunity cost, clinician authority, and evidence-chain status. The final stress cases treat recurring hospital genetics result triage as support-only; treat multilingual, cross-script, and picture-language artifacts as audit-only until source, language, encoding, simulation map, and clinical context are established; treat a real but nonclinical speculative topological theory citation as a proof-status/provenance stress case rather than medical evidence; treat Kolmogorov-Arnold Network architecture claims as model-selection evidence rather than clinical permission; treat fixed FIFO/LIFO runtime processing as unsafe when scarce budgets may be depleted in an order-sensitive way; route residual-alpha claims to orthogonal-projection confounding audit; and preserve family genomic consent boundaries when ranked preferences cycle or privacy budgets are exceeded. The monitor returns one of several runtime actions: allow, allow bounded training simulation, narrow, deny, abstain, stress-test, stress-test ordering, stress-test confounding, preserve consent boundary, or escalate to workflow diagnosis.",
    )
    add_para(
        doc,
        "For multimodal audio and telemedicine, the core runtime-assurance pattern is deliberately simple: the model may infer, but output is gated before playback or presentation.",
    )
    add_para(
        doc,
        "This can be implemented in an async or coroutine style: inference produces an intermediate object, execution suspends for an assurance check, and only a permitted representation is resumed into the user-facing output stream.",
    )
    add_para(
        doc,
        "Liquid neural networks are included only as a candidate implementation family for future streaming monitors. Their continuous-time dynamics and small recurrent state make them a natural architecture to test for ECG, audio, wearable, or telemedicine streams, but the assurance layer remains model-agnostic: liquid, transformer, classical, and rules-based models all have to pass the same endpoint, provenance, and authority gates before output is shown.",
    )
    add_code_block(
        doc,
        "async def process_audio(chunk):\n"
        "    tracks = await separator.infer(chunk)\n"
        "    safe_output = await assurance_controller.check(tracks)\n"
        "    await audio_output.play(safe_output)",
    )
    add_para(
        doc,
        "I then added a ZDD-style sparse-family audit. Each LLM response is represented as a subset of evidence features, such as verified citation, endpoint type, study design, population fit, context status, opportunity-cost status, authority boundary, and evidence-chain status. Zero-suppressed decision diagrams are appropriate because most evidence features are absent in any given claim. The audit tests whether observed stress states and permitted hard-outcome or surrogate claim forms can be represented compactly without enumerating the full power set.",
    )
    add_para(
        doc,
        "Finally, I implemented a model-selection layer. The monitor compares candidate claim models and minimizes an information-criterion-style loss that penalizes missing evidence, forbidden boundary crossings, model complexity, nested dependency burden, and authority risk. This makes the safe output a selected claim form, not merely a refusal template: the monitor can choose second-reader assistance, surrogate-only language, confirmatory-surrogate language, hard-outcome summary, stress testing, or abstention.",
    )
    add_para(
        doc,
        "The KAN stress case is included here because it is a clean model-selection example. A critical architecture review may legitimately inform when smooth basis-function models are plausible, when benchmark failures diagnose data mismatch, or when computational overhead matters. But those are method claims. They do not become medical outcome, deployment, or safety claims until the clinical evidence dimensions are also present.",
    )
    add_para(
        doc,
        "This is the agentic core of the proposal. The model may search, parse, compare, and propose, but the hidden safety function is evaluated by the runtime controller. The setup is analogous to AI Safety Gridworlds: the observed reward of producing a helpful answer is separated from the performance function that measures whether the answer respected safety, evidence, and authority boundaries.",
    )
    add_picture(doc, PROJECT / "figures" / "evidence_gate_architecture.png", 6.2, "Figure 1. Runtime assurance architecture for evidence-gated medical LLM claims.")
    gate_rows = [
        ["Citation", "Verified, local audit, course artifact, unverifiable", "Unverifiable citations are denied."],
        ["Endpoint", "Hard outcome, validated surrogate, surrogate, process, local correction", "Only hard outcomes can support hard-outcome claims."],
        ["Design", "Pragmatic RCT, RCT, validation, observational, case", "Study design caps the permission level."],
        ["Transport/context", "Population match, mismatch, temporal shift, missing material context", "Mismatch triggers abstention or stress testing."],
        ["Provenance", "Raw transcript, cross-script text, encoding state, source language", "Unverified artifacts remain audit-only."],
        ["Opportunity cost", "Addressed, unknown, unaddressed high", "High unaddressed cost blocks deployment claims."],
        ["Authority", "Clinician retained, public-health boundary needed, handoff missing", "The model never becomes the final decision-maker."],
    ]
    add_table(doc, ["Gate", "Inputs", "Runtime rule"], gate_rows, [1800, 3400, 4160])

    doc.add_heading("4. Results", level=1)
    add_para(
        doc,
        f"The harness matched all expected actions: {summary['matched_expected_count']} of {summary['scenario_count']} scenarios. Only the hard patient-outcome pragmatic RCT scenario received permission for a cautious clinician-facing outcome summary. The synthetic telemedicine and robotic-cadaver rehearsal cases were allowed as training simulation only. The surrogate detection-rate, recurring hospital genetics triage, KAN architecture, audio/captioning benchmark, SIK3/tinnitus, saliva evolutionary-genomics, consumer heart-rate/wearable physiology, and Kanazawa-style bounded evidence ambiguity cases were narrowed. The altitude/circadian and reverse-Flynn norm-drift cases were routed to generalization stress testing. The sequential-depletion case was routed to runtime ordering stress testing; the residual-alpha case was routed to confounding/projection audit; and the family genomic-consent case preserved consent boundaries rather than aggregating ranked preferences into one family authorization. The Connes/Penrose evidence-patch case abstained from transport rather than quotienting away context. The cross-script transcript, speculative topological-theory, and picture-language simulation cases abstained for provenance/proof-status audit. The fabricated-citation case was denied, the Yablo-style evidence-chain case abstained, and the CLOT case escalated to workflow diagnosis rather than blaming clinicians for low uptake.",
    )
    add_para(
        doc,
        f"The depletion-ordering experiment enumerated {depletion_summary['permutation_count']} permutations of four synthetic runtime loads with reserve {depletion_summary['reserve']}. It found the expected ordering split: increasing load order minimized raw residual-fraction cost ({depletion_summary['best_raw_cost']:.4f}), decreasing load order maximized it ({depletion_summary['worst_raw_cost']:.4f}), while decreasing load order minimized convex log-depletion risk ({depletion_summary['best_convex_log_risk']:.4f}). In the heterogeneous two-resource probe, the same pair had positive swap delta in one tail state and negative swap delta in another, demonstrating context reversal.",
    )
    add_para(
        doc,
        f"The orthogonal-projection experiment reduced the maximum absolute design-matrix inner product from {projection_summary['max_abs_design_inner_product_before']:.3f} before projection to {projection_summary['max_abs_design_inner_product_after']:.2e} after projection. This supports the confounding gate: a residual signal should not be interpreted as independent medical evidence until its relationship to known design covariates has been audited.",
    )
    add_para(
        doc,
        f"The Kelly-style exposure experiment reproduced Ferguson's favorable-game warning in abstract form: all-in play had high expected value but a ruin probability of {kelly_summary['all_in_ruin_probability']:.3f} over {kelly_summary['ferguson_example_rounds']} rounds. The monitor therefore caps claim exposure at {kelly_summary['cap']:.2f} while reliability is still uncertain.",
    )
    add_para(
        doc,
        f"The uniform-witness experiment enumerated {witness_summary['universe_size']} possible endpoint/design/population/provenance witnesses and found {witness_summary['admissible_witness_count']} admissible witnesses. A biased strongest-five display had variation distance {witness_summary['variation_distance_from_uniform']:.3f} from the uniform admissible witness distribution, showing how curated examples can make evidence look stronger than the underlying claim family.",
    )
    add_para(
        doc,
        f"The family genomic-consent experiment detected a Condorcet-style cycle over {consent_summary['option_count']} consent options across {consent_summary['relative_count']} relatives. Full raw genome reuse was blocked, trait-specific partial inclusion remained both consent-feasible and privacy-budget-feasible, and the composed epsilon for partial-plus-audit use was {consent_summary['partial_plus_audit_composed_loss']['epsilon']:.2f} under a synthetic epsilon budget of {consent_summary['runtime_budget']['epsilon']:.2f}. This is not a HIPAA revision by itself; it is a runtime test showing why family-linked genomic records need partial permissions and budget accounting.",
    )
    add_para(
        doc,
        f"The finite event-algebra experiment built a sample space of {event_summary['sample_space_size']} runtime histories and treated unsafe promotion, privacy exceedance, consent collapse, and audit trigger as measurable events. The union and complement identities held to numerical error below {max(event_summary['union_identity_error'], event_summary['complement_identity_error']):.1e}, making the planning rule explicit: histories in the audit-trigger event are routed to narrowed or audit-only output.",
    )
    add_para(
        doc,
        f"The measure-on-measures experiment defined {measure_summary['source_measure_count']} synthetic source-population measures and a meta-measure mixture over them. The maximum total-variation distance from an individual source distribution to the mixture was {measure_summary['max_source_tv_to_mixture']:.3f}; the source with highest audit-trigger probability was {measure_summary['source_with_highest_audit_probability']}. This supports source-specific testing rather than relying only on aggregate performance.",
    )
    add_para(
        doc,
        f"The Mahalanobis covariate experiment compared {mahalanobis_summary['source_count']} synthetic source-population vectors over {mahalanobis_summary['feature_count']} covariates. It flagged {len(mahalanobis_summary['flagged_sources'])} source(s) for source-specific validation at threshold {mahalanobis_summary['action_threshold']:.1f}; the largest distance was {mahalanobis_summary['max_distance']:.3f} for {mahalanobis_summary['max_distance_source']}. This turns population/covariate fit into an auditable transport gate.",
    )
    add_para(
        doc,
        f"The exact matrix counterexample probe verified a minimal 2 by 2 failure mode for a simulation-supported noncommutative AM-GM-style inequality. The rank-one projector case has right-hand side {matrix_summary['original_rhs']} and left-hand side {matrix_summary['original_lhs_exact']}, giving violation factor {matrix_summary['violation_factor_exact']}. The positive-definite perturbation with epsilon {matrix_summary['positive_definite_epsilon']} still violates the bound with margin {matrix_summary['positive_definite_margin_exact']}. In the capstone this is used only as an evaluation-design warning: reassuring simulations should be paired with small adversarial witnesses.",
    )
    add_para(
        doc,
        f"The tail/maximal-inequality probe enumerated {tail_summary['path_count']} synthetic {tail_summary['steps']}-step mean-zero runtime-drift paths. With threshold {tail_summary['threshold']} and total variance {tail_summary['total_variance']}, Kolmogorov's bound was {tail_summary['kolmogorov_bound']}; the exact crossing probability was {tail_summary['exact_crossing_probability']} and the terminal-only exceedance probability was {tail_summary['terminal_exceed_probability']}. Because some paths crossed the boundary and returned inside by the terminal time, the probe supports pathwise runtime monitoring rather than final-output-only review.",
    )
    add_para(
        doc,
        f"The Hex boundary-invariant probe enumerated {hex_summary['total_full_boards_enumerated']} full boards across sizes {', '.join(str(x) for x in hex_summary['board_sizes_enumerated'])}. It found {hex_summary['both_crossing_count']} both-crossing boards and {hex_summary['neither_crossing_count']} neither-crossing boards, so every full board had exactly one crossing. A small minimax check also found first-player wins for sizes {', '.join(str(x) for x in hex_summary['first_player_win_sizes_checked'])}. The majority-triangle subcheck enumerated {hex_summary['majority_triangle_summary']['patterns']} local triple patterns with {hex_summary['majority_triangle_summary']['tie_count']} ties. In the capstone this is a no-ambiguous-terminal-state and safe-coarse-graining check for runtime actions, not a clinical result.",
    )
    add_para(
        doc,
        f"A larger Hex scaling probe then sampled {hex_scaling_summary['sampled_full_boards']} full boards over sizes {', '.join(str(x) for x in hex_scaling_summary['crossing_board_sizes'])}. The sampled terminal ambiguity count remained {hex_scaling_summary['ambiguous_terminal_count']}, and the unbiased blue-crossing mean was {hex_scaling_summary['unbiased_blue_crossing_mean']:.3f}. However, a generic same-size local majority smoother flipped the global crossing in up to {hex_scaling_summary['max_coarse_grain_flip_rate']:.3f} of sampled boards, with the maximum at side length {hex_scaling_summary['max_coarse_grain_flip_setting']['board_size']} and blue probability {hex_scaling_summary['max_coarse_grain_flip_setting']['blue_probability']}. This is the operational lesson for genetic privacy blurring and multimodal summarization: aggregation can look locally reasonable while changing the safety boundary.",
    )
    add_para(
        doc,
        f"The ZDD-style audit used {zdd_summary['feature_universe_size']} possible evidence features. The {summary['scenario_count']} scenarios collapsed to {zdd_summary['observed_family_size']} distinct sparse feature subsets, represented by {zdd_summary['observed_zdd_node_count']} ZDD nodes versus a naive trie upper bound of {zdd_summary['naive_observed_trie_upper_bound']}. The hard-outcome family admitted the pragmatic RCT scenario and rejected surrogate cases; the surrogate/process-support family admitted surrogate-overclaim, recurring hospital genetics triage, KAN architecture, audio/captioning, SIK3/tinnitus, saliva evolutionary-genomics, and Kanazawa bounded-ambiguity cases; the bounded-training family admitted telemedicine and robotic rehearsal cases; the provenance-gap family admitted the cross-script transcript audit case, the speculative topological-theory proof-status case, and the picture-language simulation case; the orthogonal-projection family admitted the residual-alpha case; and the consent-boundary family admitted the family genomic-consent case.",
    )
    add_para(
        doc,
        f"The proof-status poset audit enumerated {proof_status_summary['states_enumerated']} states over {len(proof_status_summary['dimension_names'])} graded evidence dimensions and {proof_status_summary['cover_transitions_enumerated']} cover transitions. Only {proof_status_summary['permission_counts']['hard_outcome_allowed_with_caveats']} states permitted a hard-outcome claim, while {proof_status_summary['nonclinical_hard_state_count']} nonclinical states and {proof_status_summary['endpoint_free_promotion_count']} endpoint-free states were promoted. NIETTU remained audit-only under all source/proof-only upgrades. This makes the graded-poset point explicit: proof/source status is not exchangeable for endpoint, transport/context, or clinician-authority status.",
    )
    add_para(
        doc,
        f"The picture-language diagram audit classified {picture_summary['diagram_count']} diagrammatic or mathematical artifacts, including Jaffe-Liu picture-language simulations, an Axelrod-style repeated-game noise case, and a Bourgade-Huang loop-equation universality case. None authorized a hard patient-outcome claim. {picture_summary['transfer_reset_diagram_count']} artifacts reset validation after representation transfer, {picture_summary['virtual_state_warning_count']} required explicit virtual-state labeling, and the loop-equation case remained universality-audit-only: approximate invariant hierarchies can support portability tests, not clinical permission by resemblance.",
    )
    add_para(
        doc,
        f"The loop-equation/Gronwall stability probe then made that portability rule executable. It checked {loop_summary['required_gate_count']} required invariant gates and compared a stable multimodal transfer chain with an unstable missing-provenance chain. The stable chain ended at error {loop_summary['stable_final_error']:.4f} under budget {loop_summary['stable_budget']:.3f}; the unstable chain ended at {loop_summary['unstable_final_error']:.4f} over budget {loop_summary['unstable_budget']:.3f}. A single-entry perturbation calculation stayed below its synthetic budget, and the switching-cumulant subcheck canceled the main quadratic term while keeping rare-event and replacement-error terms inside budget. The permission therefore remained universality-audit-only rather than outcome permission.",
    )
    add_para(
        doc,
        f"The branch-factor path probe added a collision/branch-cut version of the same rule. A separated multimodal branch with R0={branch_summary['stable_r0']:.2f} preserved half-planes, positive phase margin, separation, and Volterra contraction ratio {branch_summary['stable_volterra_contraction_ratio']:.4f}. A near-collision family/population branch with R0={branch_summary['near_collision_r0']:.3f} failed the separation/contraction audit and was routed to branch-stability abstention. This operationalizes the BBGKY warning: pointwise-looking identities away from collisions do not justify near-collision genomic, population, or contact-term claims.",
    )
    add_para(
        doc,
        f"The model-selection audit compared {model_summary['candidate_model_count']} candidate claim models and matched the expected safety action in {model_summary['matched_expected_count']} of {model_summary['scenario_count']} scenarios. This result is intentionally modest: it shows that claim permission, multimodal simulation boundaries, and ordering stress tests can be made explicit before clinical deployment, not that a medical LLM has improved outcomes.",
    )
    add_picture(doc, PROJECT / "figures" / "runtime_action_counts.png", 6.2, "Figure 2. Runtime actions across synthetic stress cases.")
    add_picture(doc, PROJECT / "figures" / "endpoint_design_permission_map.png", 5.8, "Figure 3. Endpoint and design strength map. Runtime gates still apply after scoring.")
    add_picture(
        doc,
        PROJECT / "figures" / "survivorship_bias.png",
        5.8,
        "Figure 4. Survivorship-bias analogy for medical LLM evaluation: the missing failures, abstentions, and non-returning evidence chains matter. Image: Grandjean/McGeddon, Wikimedia Commons, CC BY-SA 4.0.",
    )

    add_result_matrix(doc, summary["scenario_rows"])

    doc.add_heading("5. Discussion", level=1)
    add_para(
        doc,
        "The result is a small but concrete standard for medical LLM evaluation. A model fails not only when it fabricates. It also fails when it is technically defensible at the wrong evidentiary strength, in the wrong population, at the wrong time, or with the wrong opportunity cost. This is why the capstone treats evidence as an asset and clinical claim strength as a currency. A verified citation is not freely exchangeable into any claim the model can phrase.",
    )
    add_para(
        doc,
        "For bioinformatics, the same rule applies to privacy and consent. Genetic records are not isolated facts about one patient; they are correlated with relatives, ancestral populations, and future interpretations. Differential privacy gives the right lens because it treats privacy as a measurable loss under repeated analysis, while Arrow's theorem warns against pretending that ranked family preferences always aggregate into one legitimate group decision. The proposed runtime monitor therefore preserves a graded poset of partial release states: full raw genome reuse, trait-specific inclusion, audit-only use, or no reuse may be incomparable across family members and clinical purposes.",
    )
    add_para(
        doc,
        "The framework also reconciles several course themes. The Module 4 human-baseline example supports bounded second-reader use: an LLM can help catch a possible error without becoming the authority. The Module 5 assignment becomes the central evidence-overclaiming test. The Module 6 Pinochet gesture argument becomes a context gate: an action pattern does not mean the same thing when the material, handoff, or procedural phase changes. The Module 7 SEIR simulation becomes a distribution-shift gate: a day-14 versus day-26 peak shift is exactly the kind of change a monitor should detect before allowing confident predictions.",
    )
    add_para(
        doc,
        "Yablo's paradox without self-reference supplies a useful warning about citation chains: the chain can be unstable even if no individual source is fabricated. Borrowing only metaphorically from Nakayama-style reasoning, a local citation may support a narrow claim without generating the stronger global claim that an intervention improves patient outcomes in the target population.",
    )
    add_para(
        doc,
        "Survivorship bias makes the evaluation problem even sharper. A benchmark that records only successful-looking LLM answers is like looking only at the planes that returned. The evaluation log must also preserve refusals, hidden denominator cases, translation failures, dropped citations, missing subgroups, and prompts where the correct action is silence or escalation. Runtime assurance is partly a way of making the non-survivors visible before they become invisible evidence gaps.",
    )
    add_para(
        doc,
        "The skater/Hawkes replay video included in the package is not medical evidence. It is a visual analogy for runtime assurance: delayed commands, speed envelopes, and separation constraints are monitored separately from the agent's nominal behavior. Medical LLMs need the same separation between fluent generation and safety envelopes.",
    )
    add_para(
        doc,
        "The reverse-sprinkler result supplies one more implementation analogy. Smith and colleagues show that reversing flow through a sprinkler does not merely invert the ordinary sprinkler mechanism; the geometry of the arms controls momentum flux and therefore rotation. Medical LLM workflows have a similar asymmetry. Generating a fluent clinical answer from evidence and reconstructing the supporting evidence from a fluent answer are not inverse operations unless the architecture preserves provenance, endpoint category, and boundary-crossing history.",
    )
    add_para(
        doc,
        "Audiology is a useful future stress domain because it makes surrogacy unavoidable. Genetic etiology, caption accuracy, speech perception, word recognition, device performance, and cochlear-implant use can all be clinically meaningful, but none is identical to communication, autonomy, quality of life, or Deaf cultural identity. SIK3 is therefore handled as a guarded hypothesis case: it can support hearing-associated genetic language, and tinnitus GWAS work supports polygenic and neuropsychiatric links, but neither proves tolerance for hearing loss, tinnitus adaptation, counseling benefit, or quality-of-life improvement. The film CODA is a concise cultural reminder of the same point: the relevant outcome is not simply normalization to hearing, but communication access and patient-centered flourishing.",
    )
    add_para(
        doc,
        "Altitude and circadian rhythm sharpen the transport problem. High-altitude auditory studies show that hypoxia, hypobaria, exposure duration, testing modality, and acclimatization can change measured auditory function. Circadian disruption is likewise a health modifier, while Dinacharya gives a culturally legible vocabulary for daily rhythm and routine. These are context features, not outcome proof. A runtime monitor should therefore stress-test any hearing, tinnitus, sleep, or wearable inference that moves across altitude, routine, ancestry, or care setting.",
    )
    add_para(
        doc,
        "The positive version of the same idea is telemedicine practice. In a synthetic or standardized-patient environment, audio separation, captions, multimodal replay, and LLM feedback can help clinicians rehearse webside manner, turn-taking, language access, uncertainty disclosure, and escalation boundaries. The runtime permission is training-only: the system may critique the simulated encounter, but it may not imply that the tool improves real patient outcomes until a separate patient-level evaluation supports that claim.",
    )
    add_para(
        doc,
        "The multilingual transcript case extends this to source provenance. A dub, caption file, OCR scrape, right-to-left Hebrew page, Mongolian Cyrillic devotional text, or mixed-script transcript can be meaningful in its own domain, but it is not automatically a clean medical input. Before an LLM may use such material in clinical reasoning, the runtime layer should verify source, language, encoding, speaker identity, translation status, clinical relevance, and consent boundary. This is the practical reason the paper argues against raw text as the default medical LLM input: text is already a transformation of the encounter.",
    )
    add_para(
        doc,
        "Toosarvandani's language-memory essay sharpens the healthcare version of this point. For marginalized language communities, speech is not just a carrier of propositional content; it can encode family memory, identity, discrimination, and access to care. Romani adds a concrete technical example because dialect diversity, endangered varieties, and Para-Romani contact varieties can defeat naive language labels. A runtime monitor for multimodal clinical AI should therefore treat language capture, captioning, translation, and summarization as governed transformations, not as neutral preprocessing steps.",
    )
    add_para(
        doc,
        "Robotics extends the same boundary. The useful question is not whether robots are 'better than us' in a science-fiction sense, but whether robots, synthetic organs, and humanoid standardized patients can function like cadaver-like rehearsal surfaces: repeatable, inspectable, and safe for mistakes. Literature on synthetic organs and humanoid patient robots supports training and feasibility claims, but the runtime gate should block autonomous treatment claims or patient-benefit claims until separate clinical evidence exists.",
    )
    add_para(
        doc,
        "The reverse-Flynn neuropsychology example makes the same point for psychometrics. A high-secure psychiatric population showed temporal and population-specific shifts in WAIS full-scale IQ and processing-speed profiles. A medical LLM should not treat general-population norms, old test versions, repeated IQ testing, or ordinary cognitive expectations as portable without checking the population, decade, instrument version, and clinical pathway. This is a runtime transport gate, not a generic intelligence claim.",
    )
    add_para(
        doc,
        "The Connes/Penrose analogy explains why the project uses sparse relational evidence states rather than only scalar confidence. Classical quotienting says two cases are equivalent when the local patch looks the same; runtime assurance asks which relations, contexts, and permissions survive before a claim is made. This is the same reason the capstone uses ZDD-style sparse families: the safe object is not just an answer, but a membership check in a structured family of admissible evidence states.",
    )
    add_para(
        doc,
        "The control-theory analogy is likewise bounded. While the claim policy is smooth, it resembles a Hamilton-Jacobi-Bellman problem: choose the safest action under an evidence-state value function. When a claim reaches a permission boundary, that value function becomes nonsmooth and viscosity-solution or distributional language becomes more appropriate. Korteweg-de Vries-style dispersive regularization is useful only as a metaphor for the boundary layer: the dangerous region is the oscillatory transition where surrogate claims begin to travel like outcome claims. The monitor's job is to detect that transition, not to claim that medical evidence literally obeys a wave equation.",
    )

    doc.add_heading("6. Limitations", level=1)
    add_bullets(
        doc,
        [
            "The scenarios are synthetic and derived from course artifacts and public literature, not from live patient data.",
            "The harness tests a safety architecture, not clinical effectiveness.",
            "No medical advice, diagnosis, treatment recommendation, or patient-facing product is produced.",
            "The next stage requires clinician review, a preregistered rubric, blinded LLM outputs, and eventually pragmatic trial designs if patient-outcome claims are made.",
        ],
    )

    doc.add_heading("7. Conclusion", level=1)
    add_para(
        doc,
        "Medical LLMs need evidence gates, not just better prompts. The proposed runtime-assurance framework gives a practical path: verify the citation, classify the endpoint, score the design, test population and context, account for opportunity cost, and preserve clinician authority. The most practical first use case is recurring hospital genetics triage, where verified but bounded artifacts must be routed without being converted into treatment authority. The framework is conservative by design. That conservatism is what lets the system be useful without pretending that polished language is clinical proof.",
    )

    doc.add_heading("References", level=1)
    for i, (citation, url) in enumerate(SOURCES, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(f"{i}. {citation} ")
        r.font.size = Pt(8.5)
        u = p.add_run(url)
        u.font.size = Pt(8.5)
        u.font.color.rgb = BLUE

    doc.save(path)
    return path


def create_summary_sheet(
    summary: dict,
    zdd_summary: dict,
    model_summary: dict,
    depletion_summary: dict,
    projection_summary: dict,
    kelly_summary: dict,
    witness_summary: dict,
    consent_summary: dict,
    event_summary: dict,
    measure_summary: dict,
    mahalanobis_summary: dict,
) -> Path:
    path = PACKAGE / "summary" / "Module_14_Summary_Sheet_Evidence_Gated_Medical_LLMs_Ravi_Bajaj.docx"
    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    style_doc(
        doc,
        "Summary Sheet",
        "Conference metadata, abstract, AI-use disclosure, and package contents",
        "Summary Sheet",
    )

    doc.add_heading("Conference / Symposium", level=1)
    rows = [
        ["Name", "Machine Learning for Health Symposium (ML4H 2026), Proceedings Track"],
        ["Conference URL", "https://ml4h.ahli.cc/"],
        ["Author instructions / CFP", "https://ml4h.ahli.cc/submit/call-for-papers/"],
        ["Submission deadline", "September 10, 2026, 11:59 PM AoE"],
        ["Author notification", "October 22, 2026"],
        ["Camera-ready deadline", "November 7, 2026 (tentative)"],
        ["Event dates", "December 6-7, 2026"],
        ["Location", "Sydney, Australia"],
        ["Submission type", "Proceedings paper: up to 8 pages at submission, excluding references and appendices; accepted papers are published in PMLR."],
    ]
    add_table(doc, ["Field", "Value"], rows, [2600, 6760])

    doc.add_heading("Final Paper Title and Abstract", level=1)
    add_para(doc, "Title: Surrogate-Aware Runtime Assurance for Agentic Medical LLMs")
    add_para(
        doc,
        "Abstract: Medical large language models (LLMs) are often evaluated for hallucination and citation accuracy, but agentic clinical workflows introduce a subtler failure: evidence overclaiming by systems that retrieve, summarize, plan, and draft clinical-sounding actions. This paper proposes a runtime-assurance framework that gates LLM outputs and proposed actions by citation validity, endpoint strength, study design, population fit, actionability, opportunity cost, clinician authority, source provenance, runtime budget order, confounding, proportional exposure, witness selection, privacy loss, and consent aggregation. A synthetic stress-test set distinguishes hard outcomes from surrogate/process endpoints, flags unverifiable citations, treats low clinician uptake as possible safety filtering, and requires abstention when context is missing. Added probes cover recurring hospital genetics triage, including pharmacogenomic alerts, newborn-screening follow-up, ACMG secondary findings, hereditary cancer/FH/Lynch flags, VUS handling, and family-linked privacy boundaries, plus multimodal telemedicine, wearable physiology, cross-script provenance, sequential depletion, orthogonal projection, Kelly-style exposure limits, uniform witness sampling, and family genomic consent under differential-privacy-style budget accounting. It does not validate a clinical product, but provides a testable architecture for safer clinician-facing agentic medical LLMs.",
    )

    doc.add_heading("Generative AI Disclosure", level=1)
    add_para(
        doc,
        "Generative AI was used to assist with brainstorming, literature organization, code scaffolding, document drafting, revision, and formatting. The author reviewed, directed, and revised the conceptual framing; selected the capstone direction; provided the course materials, prior assignment insights, and instructor feedback; and remains responsible for all claims, citations, code, and final submissions. No real patient data were used. The synthetic stress tests do not provide medical advice, diagnosis, or treatment recommendations.",
    )

    doc.add_heading("Optional Data, Code, and Tools", level=1)
    add_bullets(
        doc,
        [
            f"Deterministic evidence-gate stress-test harness with {summary['scenario_count']} synthetic scenarios.",
            f"ZDD-style sparse-family audit over {zdd_summary['feature_universe_size']} evidence features.",
            f"Model-selection audit over {model_summary['candidate_model_count']} candidate claim forms.",
            "Stage 1 probes: depletion ordering, orthogonal projection, Kelly exposure, uniform witnesses, consent/privacy budget, finite event algebra, measure-on-measures, and Mahalanobis covariate distance.",
            f"Consent audit blocks full raw genomic release while preserving partial inclusion; Mahalanobis flags {', '.join(mahalanobis_summary['flagged_sources'])} for source-specific validation.",
            "CSV/JSON results, reproducible Python tests, generated figures, and one demo video are included; all are synthetic or analogy-only, not clinical evidence.",
        ],
    )

    doc.save(path)
    return path


def create_technical_supplement(
    summary: dict,
    zdd_summary: dict,
    model_summary: dict,
    depletion_summary: dict,
    projection_summary: dict,
    kelly_summary: dict,
    witness_summary: dict,
    consent_summary: dict,
    event_summary: dict,
    measure_summary: dict,
    mahalanobis_summary: dict,
    matrix_summary: dict,
    tail_summary: dict,
    hex_summary: dict,
    hex_scaling_summary: dict,
    proof_status_summary: dict,
    picture_summary: dict,
    loop_summary: dict,
    branch_summary: dict,
) -> Path:
    path = PACKAGE / "supplement" / "Module_14_Technical_Supplement_Evidence_Gated_Medical_LLMs_Ravi_Bajaj.docx"
    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    style_doc(
        doc,
        "Technical Supplement",
        "Stress-test records, runtime invariants, and reproducibility notes",
        "Technical Supplement",
    )

    doc.add_heading("Safety Boundary", level=1)
    add_bullets(
        doc,
        [
            "Synthetic-only evaluation artifact.",
            "No real patient data, no clinical deployment, no diagnosis, and no treatment recommendation.",
            "The system evaluates claim permission levels, not patient care.",
            "Human clinician or public-health authority remains outside the LLM and is explicitly represented in each scenario.",
        ],
    )

    doc.add_heading("Runtime Invariant", level=1)
    add_para(doc, summary["safety_invariant"])
    add_para(doc, "For multimodal or streaming systems, the assurance controller sits between inference and output:")
    add_code_block(
        doc,
        "async def process_audio(chunk):\n"
        "    tracks = await separator.infer(chunk)\n"
        "    safe_output = await assurance_controller.check(tracks)\n"
        "    await audio_output.play(safe_output)",
    )
    add_para(
        doc,
        "Coroutine-style control flow is useful here because the monitor can suspend output at the boundary where evidence, provenance, or clinician authority has not yet been established.",
    )
    add_para(
        doc,
        "Liquid neural networks are a plausible future architecture for the streaming inference block because continuous-time recurrent dynamics are well matched to ECG, wearable, audio, and telemedicine streams. This supplement treats them as an implementation candidate only. The assurance contract is intentionally independent of model class, so a liquid neural network would still have to emit auditable features that can be checked by the same endpoint, provenance, and authority gates.",
    )

    doc.add_heading("Hospital Genetics Scope", level=1)
    add_bullets(
        doc,
        [
            "Pharmacogenomic alerts, especially medication-gene pairs such as CYP2C19 and clopidogrel.",
            "Positive newborn-screening follow-up routed through ACT sheets, algorithms, and local/state newborn-screening workflows.",
            "ACMG secondary findings and CDC Tier 1 applications: hereditary breast and ovarian cancer, Lynch syndrome, and familial hypercholesterolemia.",
            "Variants of uncertain significance, where the permitted LLM role is documentation/routing/explanation rather than clinical interpretation beyond the report.",
            "Family-linked privacy and partial inclusion decisions, where one person's genome can reveal information about relatives.",
        ],
    )

    doc.add_page_break()
    doc.add_heading("Stress-Test Matrix", level=1)
    rows = []
    for row in summary["scenario_rows"]:
        rows.append(
            [
                row["scenario_id"],
                row["label"],
                ACTION_SHORT.get(row["actual_action"], row["actual_action"].replace("_", " ").lower()),
                row["blocked_by"] or "none",
            ]
        )
    add_para(doc, "The full scenario records, including requested claims and permission levels, are provided in results/evidence_gate_stress_test_rows.csv.")
    midpoint = (len(rows) + 1) // 2
    add_table(doc, ["ID", "Scenario", "Runtime action", "Blocked by"], rows[:midpoint], [650, 4300, 2300, 2110])
    doc.add_page_break()
    add_table(doc, ["ID", "Scenario", "Runtime action", "Blocked by"], rows[midpoint:], [650, 4300, 2300, 2110])

    doc.add_heading("Decision Counts", level=1)
    count_rows = [[action, str(count)] for action, count in summary["decision_counts"].items()]
    add_table(doc, ["Runtime action", "Count", "Runtime action", "Count"], pair_rows(count_rows), [3600, 900, 3600, 900])

    doc.add_page_break()
    doc.add_heading("ZDD Sparse-Family Audit", level=1)
    add_para(
        doc,
        "The ZDD audit treats every LLM claim state as a sparse subset of possible evidence features. This is a data-structure version of the capstone's main argument: the monitor checks membership in a pre-specified family of safe claim forms rather than relying only on scalar scores.",
    )
    zdd_rows = [
        ["Feature universe", str(zdd_summary["feature_universe_size"])],
        ["Observed distinct feature subsets", str(zdd_summary["observed_family_size"])],
        ["Observed ZDD node count", str(zdd_summary["observed_zdd_node_count"])],
        ["Naive trie upper bound", str(zdd_summary["naive_observed_trie_upper_bound"])],
        ["All observed scenarios recognized", str(zdd_summary["observed_membership_all_true"])],
        ["Hard-outcome family admits S03", str(zdd_summary["s03_hard_outcome_allowed"])],
        ["Hard-outcome family rejects S02", str(zdd_summary["s02_hard_outcome_allowed"])],
        ["Surrogate family admits S02", str(zdd_summary["s02_surrogate_allowed"])],
        ["Surrogate family admits S13", str(zdd_summary["s13_surrogate_allowed"])],
        ["Surrogate family admits S14", str(zdd_summary["s14_surrogate_allowed"])],
        ["Surrogate family admits S16", str(zdd_summary["s16_surrogate_allowed"])],
        ["Surrogate family admits S22", str(zdd_summary["s22_surrogate_allowed"])],
        ["Surrogate family admits S24", str(zdd_summary["s24_surrogate_allowed"])],
        ["Surrogate family admits S27", str(zdd_summary["s27_surrogate_allowed"])],
        ["Surrogate family admits S29", str(zdd_summary["s29_surrogate_allowed"])],
        ["Surrogate family rejects S17", str(zdd_summary["s17_surrogate_allowed"])],
        ["Surrogate family rejects S19", str(zdd_summary["s19_surrogate_allowed"])],
        ["Training family admits S15", str(zdd_summary["s15_bounded_training_allowed"])],
        ["Training family rejects S14", str(zdd_summary["s14_bounded_training_allowed"])],
        ["Training family admits S18", str(zdd_summary["s18_bounded_training_allowed"])],
        ["Runtime-ordering family admits S23", str(zdd_summary["s23_runtime_ordering_allowed"])],
        ["Runtime-ordering family rejects S15", str(zdd_summary["s15_runtime_ordering_allowed"])],
        ["Orthogonal-projection family admits S25", str(zdd_summary["s25_orthogonal_projection_allowed"])],
        ["Orthogonal-projection family rejects S15", str(zdd_summary["s15_orthogonal_projection_allowed"])],
        ["Consent-boundary family admits S26", str(zdd_summary["s26_consent_boundary_allowed"])],
        ["Consent-boundary family rejects S15", str(zdd_summary["s15_consent_boundary_allowed"])],
        ["Provenance family admits S21", str(zdd_summary["s21_provenance_gap_allowed"])],
        ["Provenance family admits S28", str(zdd_summary["s28_provenance_gap_allowed"])],
        ["Provenance family admits S30", str(zdd_summary["s30_provenance_gap_allowed"])],
    ]
    add_table(doc, ["Audit item", "Value"], zdd_rows, [5100, 2200])

    doc.add_page_break()
    doc.add_heading("Model-Selection Audit", level=1)
    add_para(
        doc,
        "The model-selection layer makes the claim grammar operational. Instead of mapping every unsafe state to a generic refusal, it compares candidate claim models and selects the least unsafe action compatible with the evidence. The loss function is information-criterion-like: missing required evidence, forbidden boundary crossings, complexity, nested dependency burden, and authority risk are all penalized.",
    )
    model_rows = [
        ["Scenario count", str(model_summary["scenario_count"])],
        ["Candidate claim models", str(model_summary["candidate_model_count"])],
        ["Matched expected actions", str(model_summary["matched_expected_count"])],
        ["All expected actions matched", str(model_summary["all_matched_expected"])],
        ["Selection rule", model_summary["model_selection_rule"]],
        ["FIFO/LIFO policy", model_summary["fifo_lifo_policy"]],
    ]
    add_table(doc, ["Audit item", "Value"], model_rows, [2500, 6500])

    doc.add_page_break()
    doc.add_heading("Proof-Status / Provenance Poset Probe", level=1)
    add_para(
        doc,
        "This finite graded-poset audit separates source existence and proof status from clinical evidence. It tests the exact issue raised by mathematical/theoretical sources: an impressive or real source object can improve provenance confidence without becoming a clinical endpoint, target-context, or authority-boundary credential.",
    )
    proof_rows = [
        ["Clinical status", proof_status_summary["clinical_status"]],
        ["Evidence dimensions", ", ".join(proof_status_summary["dimension_names"])],
        ["States enumerated", str(proof_status_summary["states_enumerated"])],
        ["Cover transitions enumerated", str(proof_status_summary["cover_transitions_enumerated"])],
        [
            "Hard-outcome permitted states",
            str(proof_status_summary["permission_counts"]["hard_outcome_allowed_with_caveats"]),
        ],
        ["Nonclinical hard-state violations", str(proof_status_summary["nonclinical_hard_state_count"])],
        ["Endpoint-free promotions", str(proof_status_summary["endpoint_free_promotion_count"])],
        ["Method-only hard-state violations", str(proof_status_summary["method_only_hard_state_count"])],
        ["NIETTU permission", proof_status_summary["niettu_permission"]],
        [
            "NIETTU source/proof-upgrade permissions",
            ", ".join(proof_status_summary["niettu_source_proof_upgrade_permissions"]),
        ],
    ]
    add_table(doc, ["Audit item", "Value"], proof_rows, [3300, 5700])
    add_picture(
        doc,
        PROJECT / "figures" / "proof_status_poset_permissions.png",
        6.4,
        "Figure S1. Proof-status/provenance poset audit. Source and proof upgrades do not substitute for clinical endpoint, context, and authority gates.",
    )

    doc.add_heading("Sequential Depletion Ordering Experiment", level=1)
    add_para(
        doc,
        "The depletion experiment is synthetic and nonclinical. It tests whether scarce runtime resources can be treated as if order does not matter. The answer is no: the cost criterion determines whether small-first or large-first is preferred, and heterogeneous resource vectors can reverse preferred order across residual states.",
    )
    depletion_rows = [
        ["Synthetic loads", ", ".join(str(x) for x in depletion_summary["loads"])],
        ["Reserve", str(depletion_summary["reserve"])],
        ["Permutations enumerated", str(depletion_summary["permutation_count"])],
        ["Increasing minimizes raw cost", str(depletion_summary["increasing_minimizes_raw_cost"])],
        ["Decreasing maximizes raw cost", str(depletion_summary["decreasing_maximizes_raw_cost"])],
        ["Decreasing minimizes convex log risk", str(depletion_summary["decreasing_minimizes_convex_log_risk"])],
        ["Increasing maximizes convex log risk", str(depletion_summary["increasing_maximizes_convex_log_risk"])],
        ["Context reversal observed", str(depletion_summary["context_reversal_observed"])],
        ["Low-tail swap delta", f"{depletion_summary['low_tail_delta_cost_a_then_b_minus_b_then_a']:.6f}"],
        ["Mixed-tail swap delta", f"{depletion_summary['mixed_tail_delta_cost_a_then_b_minus_b_then_a']:.6f}"],
    ]
    add_table(doc, ["Experiment item", "Value"], depletion_rows, [4200, 4200])

    doc.add_heading("Orthogonal Projection / Confounding Experiment", level=1)
    add_para(
        doc,
        "The projection experiment is synthetic and nonclinical. It tests whether a latent residual component can be treated as independent evidence before checking its relationship to known design covariates. The runtime action is projection/confounding audit only.",
    )
    projection_rows = [
        ["Synthetic sample count", str(projection_summary["sample_count"])],
        ["Design columns", str(projection_summary["design_columns"])],
        ["Alpha norm before projection", f"{projection_summary['alpha_norm_before_projection']:.6f}"],
        ["Alpha norm after projection", f"{projection_summary['alpha_norm_after_projection']:.6f}"],
        ["Max abs design inner product before", f"{projection_summary['max_abs_design_inner_product_before']:.6f}"],
        ["Max abs design inner product after", f"{projection_summary['max_abs_design_inner_product_after']:.2e}"],
        ["Orthogonality passed", str(projection_summary["orthogonality_passed"])],
    ]
    add_table(doc, ["Experiment item", "Value"], projection_rows, [4200, 4200])

    doc.add_heading("Kelly-Style Runtime Exposure Experiment", level=1)
    add_para(
        doc,
        "The Kelly experiment is a bounded analogy for claim exposure, not a betting recommendation. It shows why a favorable-looking signal should still spend only a capped fraction of clinical-claim authority while endpoint, population, provenance, and opportunity-cost gates remain uncertain.",
    )
    kelly_rows = [
        ["Favorable-game p", f"{kelly_summary['ferguson_example_probability']:.3f}"],
        ["Rounds", str(kelly_summary["ferguson_example_rounds"])],
        ["All-in expected factor", f"{kelly_summary['all_in_expected_factor']:.3f}"],
        ["All-in ruin probability", f"{kelly_summary['all_in_ruin_probability']:.3f}"],
        ["Kelly fraction", f"{kelly_summary['kelly_fraction_for_two_thirds']:.3f}"],
        ["Exposure cap", f"{kelly_summary['cap']:.2f}"],
    ]
    add_table(doc, ["Experiment item", "Value"], kelly_rows, [4200, 4200])

    doc.add_heading("Uniform Witness Sampling Experiment", level=1)
    add_para(
        doc,
        "The witness experiment is a synthetic audit for cherry-picking. Evidence witnesses are endpoint/design/population/provenance tuples. A biased display of only the strongest-looking witnesses can make a claim family look more outcome-ready than it is.",
    )
    witness_rows = [
        ["Witness universe size", str(witness_summary["universe_size"])],
        ["Admissible witness count", str(witness_summary["admissible_witness_count"])],
        ["Biased top-witness count", str(witness_summary["biased_top_witness_count"])],
        ["Variation distance from uniform", f"{witness_summary['variation_distance_from_uniform']:.3f}"],
        ["Biased hard-outcome fraction", f"{witness_summary['biased_hard_outcome_fraction']:.3f}"],
        ["Uniform hard-outcome fraction", f"{witness_summary['uniform_hard_outcome_fraction']:.3f}"],
    ]
    add_table(doc, ["Experiment item", "Value"], witness_rows, [4200, 4200])

    doc.add_heading("Family Genomic Consent / Privacy-Budget Experiment", level=1)
    add_para(
        doc,
        "The consent experiment is a synthetic governance test for family-linked genomic data. It detects a ranked-preference cycle, preserves individual consent caps, and applies differential-privacy-style runtime budget accounting before permitting any release state.",
    )
    consent_rows = [
        ["Consent options", str(consent_summary["option_count"])],
        ["Relatives", str(consent_summary["relative_count"])],
        ["Condorcet-style cycle detected", str(consent_summary["condorcet_cycle_detected"])],
        ["Full raw genome release runtime-safe", str(consent_summary["full_raw_release_runtime_safe"])],
        ["Trait-specific partial inclusion runtime-safe", str(consent_summary["partial_inclusion_runtime_safe"])],
        ["Audit-only runtime-safe", str(consent_summary["audit_only_runtime_safe"])],
        ["Runtime epsilon budget", f"{consent_summary['runtime_budget']['epsilon']:.2f}"],
        ["Partial-plus-audit composed epsilon", f"{consent_summary['partial_plus_audit_composed_loss']['epsilon']:.2f}"],
        ["Three-relative partial group epsilon bound", f"{consent_summary['group_privacy_bound_for_partial_three_relatives']['epsilon']:.2f}"],
    ]
    add_table(doc, ["Experiment item", "Value"], consent_rows, [4200, 4200])
    add_para(
        doc,
        "This experiment does not claim compliance with HIPAA or any NHS policy. Its narrower purpose is to make a HIPAA-adjacent design question testable: if a future rule permits privacy-budgeted genomic analysis, the runtime system should explicitly log neighboring-record granularity, group/family effects, composition, and post-processing boundaries before a medical LLM can use the released features.",
    )

    doc.add_heading("Finite Event Algebra Experiment", level=1)
    add_para(
        doc,
        "This experiment makes the measure-theory term event concrete. In the finite sample space, every subset of runtime histories is measurable. Unsafe promotion, privacy exceedance, consent collapse, and audit trigger are events; the runtime plan maps those events to blocked, narrowed, or audit-only actions.",
    )
    event_rows = [
        ["Runtime histories", str(event_summary["sample_space_size"])],
        ["Unsafe-promotion event size", str(event_summary["unsafe_promotion_event_size"])],
        ["Audit-trigger event size", str(event_summary["audit_trigger_event_size"])],
        ["Safe-stop event size", str(event_summary["safe_stop_event_size"])],
        ["Unsafe-promotion probability", f"{event_summary['unsafe_promotion_probability']:.3f}"],
        ["Audit-trigger probability", f"{event_summary['audit_trigger_probability']:.3f}"],
        ["Safe-stop probability", f"{event_summary['safe_stop_probability']:.3f}"],
        ["Union identity error", f"{event_summary['union_identity_error']:.2e}"],
        ["Complement identity error", f"{event_summary['complement_identity_error']:.2e}"],
        ["Monotonicity holds", str(event_summary["monotonicity_holds"])],
    ]
    add_table(doc, ["Experiment item", "Value"], event_rows, [4200, 4200])

    doc.add_heading("Measure-on-Measures Source-Robustness Experiment", level=1)
    add_para(
        doc,
        "Each synthetic source population induces a probability measure over runtime histories. A meta-measure over those source measures represents uncertainty about which source distribution is active. The mixture is useful, but it must not hide the highest-risk source distribution.",
    )
    measure_rows = [
        ["History-space size", str(measure_summary["history_space_size"])],
        ["Source measures", str(measure_summary["source_measure_count"])],
        ["Meta-measure total mass", f"{measure_summary['meta_measure_total_mass']:.3f}"],
        ["Mixture unsafe-promotion probability", f"{measure_summary['mixture_summary']['unsafe_promotion_probability']:.3f}"],
        ["Mixture audit-trigger probability", f"{measure_summary['mixture_summary']['audit_trigger_probability']:.3f}"],
        ["Max source TV to mixture", f"{measure_summary['max_source_tv_to_mixture']:.3f}"],
        ["Highest audit-probability source", measure_summary["source_with_highest_audit_probability"]],
    ]
    add_table(doc, ["Experiment item", "Value"], measure_rows, [4200, 4200])

    doc.add_heading("Probabilistic Programming Implementation Note", level=1)
    add_para(
        doc,
        "Figaro would be a plausible future implementation substrate for these source and event-history models because it supports rich probabilistic model definitions and reusable reasoning algorithms over evidence. The submitted package intentionally does not depend on Figaro; the Stage 1 artifact remains small Python code so the instructors can replay the logic without installing a Scala probabilistic-programming stack.",
    )

    doc.add_heading("Mahalanobis Covariate-Transport Experiment", level=1)
    add_para(
        doc,
        "This experiment converts source-population fit into a covariance-adjusted distance. It is a transport gate: sources far from the reference covariate structure require source-specific validation before an aggregate claim is allowed.",
    )
    mahalanobis_rows = [
        ["Covariates", str(mahalanobis_summary["feature_count"])],
        ["Sources", str(mahalanobis_summary["source_count"])],
        ["Action threshold", f"{mahalanobis_summary['action_threshold']:.2f}"],
        ["Flagged sources", ", ".join(mahalanobis_summary["flagged_sources"])],
        ["Max-distance source", mahalanobis_summary["max_distance_source"]],
        ["Max distance", f"{mahalanobis_summary['max_distance']:.3f}"],
    ]
    add_table(doc, ["Experiment item", "Value"], mahalanobis_rows, [4200, 4200])

    doc.add_heading("Exact Matrix Counterexample Probe", level=1)
    add_para(
        doc,
        "This is a mathematical evaluation analogy only. It checks that a small exact adversarial witness can defeat a conjecture that had simulation support, reinforcing the capstone's insistence on counterexample search for medical LLM evaluation.",
    )
    matrix_rows = [
        ["Clinical status", matrix_summary["clinical_status"]],
        ["Dimension", str(matrix_summary["dimension"])],
        ["Rank-one projectors", str(matrix_summary["rank_one_projectors"])],
        ["Dimension-one scalar case satisfies bound", str(matrix_summary["dimension_one_scalar_case_satisfies_bound"])],
        ["Original RHS", matrix_summary["original_rhs"]],
        ["Original LHS", matrix_summary["original_lhs_exact"]],
        ["Violation factor", matrix_summary["violation_factor_exact"]],
        ["Positive-definite epsilon", matrix_summary["positive_definite_epsilon"]],
        ["Positive-definite RHS", matrix_summary["positive_definite_rhs"]],
        ["Positive-definite LHS", matrix_summary["positive_definite_lhs_exact"]],
        ["Positive-definite margin", matrix_summary["positive_definite_margin_exact"]],
    ]
    add_table(doc, ["Experiment item", "Value"], matrix_rows, [4200, 4200])

    doc.add_heading("Tail / Maximal-Inequality Runtime Probe", level=1)
    add_para(
        doc,
        "This exact finite enumeration translates Sheffield's maximal-inequality lecture into a runtime-monitoring test. The key distinction is pathwise crossing versus terminal-state inspection.",
    )
    tail_rows = [
        ["Clinical status", tail_summary["clinical_status"]],
        ["Steps", str(tail_summary["steps"])],
        ["Paths enumerated", str(tail_summary["path_count"])],
        ["Increment model", tail_summary["increment_model"]],
        ["Total variance", str(tail_summary["total_variance"])],
        ["Threshold", str(tail_summary["threshold"])],
        ["Kolmogorov bound", tail_summary["kolmogorov_bound"]],
        ["Exact crossing probability", tail_summary["exact_crossing_probability"]],
        ["Terminal exceedance probability", tail_summary["terminal_exceed_probability"]],
        ["Returned inside after crossing probability", tail_summary["returned_inside_after_crossing_probability"]],
        ["Bound holds", str(tail_summary["bound_holds"])],
    ]
    add_table(doc, ["Experiment item", "Value"], tail_rows, [4200, 4200])

    doc.add_heading("Hex / Y Boundary and Coarse-Graining Probe", level=1)
    add_para(
        doc,
        "This finite audit uses Hex as a no-ambiguous-terminal-state analogy and the Y majority-triangle reduction as a safe-coarse-graining analogy. The clinical point is that blurring or aggregation should preserve the safety-relevant boundary being audited.",
    )
    majority = hex_summary["majority_triangle_summary"]
    hex_rows = [
        ["Clinical status", hex_summary["clinical_status"]],
        ["Board sizes enumerated", ", ".join(str(x) for x in hex_summary["board_sizes_enumerated"])],
        ["Full boards enumerated", str(hex_summary["total_full_boards_enumerated"])],
        ["Both-crossing boards", str(hex_summary["both_crossing_count"])],
        ["Neither-crossing boards", str(hex_summary["neither_crossing_count"])],
        ["Exactly one crossing for all full boards", str(hex_summary["exactly_one_crossing_all_full_boards"])],
        ["First-player win sizes checked", ", ".join(str(x) for x in hex_summary["first_player_win_sizes_checked"])],
        ["First player wins all checked", str(hex_summary["first_player_wins_all_checked"])],
        ["Majority-triangle patterns", str(majority["patterns"])],
        ["Majority-triangle ties", str(majority["tie_count"])],
        ["Blue/yellow majority pattern counts", f"{majority['blue_majority_patterns']} / {majority['yellow_majority_patterns']}"],
    ]
    add_table(doc, ["Experiment item", "Value"], hex_rows, [4200, 4200])

    doc.add_heading("Hex Scaling and Coarse-Graining Stress Probe", level=1)
    add_para(
        doc,
        "This Monte Carlo extension samples larger Hex boards and then tests a deliberately generic local majority smoother. The point is not to discover Hex theory; it is to show that a local blurring rule can change a global runtime boundary unless the reduction is certified to preserve that boundary.",
    )
    hex_scaling_rows = [
        ["Clinical status", hex_scaling_summary["clinical_status"]],
        ["GPU/runtime status", hex_scaling_summary["gpu_status"]],
        ["Crossing samples per setting", str(hex_scaling_summary["crossing_samples_per_setting"])],
        ["Smoothing samples per setting", str(hex_scaling_summary["smoothing_samples_per_setting"])],
        ["Crossing board sizes", ", ".join(str(x) for x in hex_scaling_summary["crossing_board_sizes"])],
        ["Smoothing board sizes", ", ".join(str(x) for x in hex_scaling_summary["smoothing_board_sizes"])],
        ["Sampled full boards", str(hex_scaling_summary["sampled_full_boards"])],
        ["Sampled smoothed boards", str(hex_scaling_summary["sampled_smoothed_boards"])],
        ["Ambiguous terminal count", str(hex_scaling_summary["ambiguous_terminal_count"])],
        ["Unbiased blue-crossing mean", f"{hex_scaling_summary['unbiased_blue_crossing_mean']:.3f}"],
        ["Max generic-smoothing flip rate", f"{hex_scaling_summary['max_coarse_grain_flip_rate']:.3f}"],
        [
            "Max flip setting",
            f"n={hex_scaling_summary['max_coarse_grain_flip_setting']['board_size']}, p={hex_scaling_summary['max_coarse_grain_flip_setting']['blue_probability']}",
        ],
    ]
    add_table(doc, ["Experiment item", "Value"], hex_scaling_rows, [4200, 4200])
    add_picture(
        doc,
        PROJECT / "figures" / "hex_scaling_coarse_grain.png",
        6.5,
        "Figure S1. Larger random Hex boards preserve no both/none terminal ambiguity in samples, while generic local smoothing can flip the global crossing.",
    )

    doc.add_heading("Event Ordering", level=1)
    order = zdd_summary["event_ordering_policy"]
    add_bullets(
        doc,
        [
            f"Audit log: {order['audit_log']}. Evidence events should be replayable in arrival order.",
            f"ZDD construction: {order['zdd_construction']}. Recursive construction may use stack-like backtracking.",
            f"ZDD membership: {order['zdd_membership']}. Membership follows the fixed feature order, not event arrival order.",
            f"Runtime safety decision: {order['runtime_safety_decision']}. Boundary crossings get safety priority over ordinary FIFO processing.",
            "Sequential depletion experiment: FIFO and LIFO are replay/search conventions, not universal policies for consuming scarce clinician-attention, audit, compute, or provenance budgets.",
        ],
    )

    doc.add_heading("Reproducibility", level=1)
    add_numbered(
        doc,
        [
            "Run run_evidence_gate_stress_test.py to regenerate JSON, CSV, and figures.",
            "Run zdd_sparse_claim_family.py to regenerate the ZDD sparse-family audit.",
            "Run sequential_depletion_verification.py to independently check the residual-depletion manuscript identities and examples.",
            "Run orthogonal_projection_experiment.py, kelly_runtime_budget_experiment.py, uniform_witness_sampling_experiment.py, and consent_aggregation_experiment.py to regenerate the added Stage 1 probes.",
            "Run pytest on the evidence-gate and ZDD test files to verify invariants.",
            "Review evidence_gate_stress_test_rows.csv for scenario-level decisions.",
            "Use the package manifest to confirm file hashes after transfer.",
        ],
    )

    doc.add_page_break()
    doc.add_heading("Notes on Gold/Kanazawa Analogy", level=1)
    add_para(
        doc,
        "This supplement does not claim that clinical evidence literally is a categorial grammar. The analogy is methodological. Gold-style learning emphasizes what can be identified from an input stream. Kanazawa's bounded k-valued classes suggest an ambiguity budget. The runtime monitor applies the same discipline to evidence: a citation may have a bounded set of admissible evidence categories, and claim composition is blocked until that category assignment is explicit.",
    )
    add_para(
        doc,
        "Finite elasticity is used only in its technical learnability sense. It is not the elasticity of rubber-band knot diagrams, even though that geometric metaphor can be helpful for intuition. Chatterjee's superconcentration/subroughness result is cited only as a probability analogy for why a monitor should track hidden variability rather than accepting a smooth aggregate claim.",
    )

    doc.add_heading("Connes/Penrose Quotient Analogy", level=1)
    add_para(
        doc,
        "Connes's Penrose-tiling example is used as a precise metaphor for context loss. If local patches recur across many global tilings, a naive quotient can erase the very structure one wanted to study. In medical LLM evaluation, repeated evidence patches - the same phrase, citation, biomarker, score, or workflow result - should not be quotient-collapsed into one universal clinical claim. The monitor preserves relational provenance through feature subsets and explicit transport gates.",
    )

    doc.add_heading("Nonlocal, Inpainting, and Coding Probe", level=1)
    add_para(
        doc,
        "The nonlocal-modeling, image-inpainting, and coding-theory references are not clinical validation sources. They supply boundary vocabulary. Nonlocal modeling says a local record may depend on distant people, prior time points, or family/social context. Inpainting says reconstructed missing content must remain labeled as reconstructed. Coding theory says transformed messages must remain decodable by the intended receiver. In the capstone, the runtime monitor applies these ideas to multimodal medical inputs, differential-privacy-style blurring, captioning, summarization, and role-specific evidence access.",
    )

    doc.add_heading("Picture-Language / Multimodal Simulation Probe", level=1)
    add_para(
        doc,
        "Jaffe and Liu's picture-language program distinguishes a formal picture language L, a target reality R, and a simulation map S from one to the other. This is the cleanest mathematical framing for multimodal medical inputs. A wound photograph, waveform, imaging frame, support-surface diagram, dashboard screenshot, or captioned video may be a useful artifact in L, but the LLM must not treat it as clinical evidence until the simulation map to R is explicit, provenance-preserved, consented, and endpoint-validated. The S30 runtime action is therefore provenance audit only.",
    )
    add_para(
        doc,
        f"The executable diagram audit extends that rule to {picture_summary['diagram_count']} artifacts. It adds Axelrod-style noisy repeated cooperation as a consent/adoption audit case and Bourgade-Huang loop-equation characterization as a universality audit case: similarity to a known pattern is not enough; the relevant invariant hierarchy must be verified before transport.",
    )
    diagram_rows = [
        [
            row["diagram_id"].replace("_", " "),
            row["source_family"].split(" program")[0],
            row["required_gate"],
            row["permission"].replace("_", " ").lower(),
        ]
        for row in picture_summary["rows"]
    ]
    add_table(doc, ["Artifact", "Source family", "Gate", "Runtime permission"], diagram_rows, [2100, 2800, 2350, 2110])
    add_picture(
        doc,
        PROJECT / "figures" / "picture_language_runtime_map.png",
        6.5,
        "Figure S2. Picture-language runtime map: multimodal artifacts remain audit/stress objects until the simulation map and runtime gates pass.",
    )

    doc.add_heading("Loop-Equation / Gronwall Stability Probe", level=1)
    add_para(
        doc,
        "The Bourgade-Huang excerpt is used as a methodological template: universality is earned by approximate loop equations, local laws, perturbation stability, and error control, not by resemblance to a known ensemble. In runtime-assurance terms, a medical LLM claim can be transported across modalities, populations, or sparse networks only when the relevant invariant gates remain small under Gronwall-style propagation.",
    )
    loop_rows = [
        ["Required gates", str(loop_summary["required_gate_count"])],
        ["Stable chain final error", f"{loop_summary['stable_final_error']:.4f} <= {loop_summary['stable_budget']:.3f}"],
        ["Unstable chain final error", f"{loop_summary['unstable_final_error']:.4f} > {loop_summary['unstable_budget']:.3f}"],
        ["Single-entry bound", f"{loop_summary['single_entry']['single_entry_bound']:.4f} < {loop_summary['single_entry']['single_entry_budget']:.3f}"],
        ["Main cumulant cancellation", f"net={loop_summary['cumulant_budget']['main_cancellation']['net']:.1f}; passes={loop_summary['cumulant_budget']['main_cancellation_passes']}"],
        ["Rare-event expectation", f"{loop_summary['cumulant_budget']['bad_event_expectation_bound']:.2e}; negligible={loop_summary['cumulant_budget']['bad_event_negligible']}"],
        ["Replacement error scale", f"{loop_summary['cumulant_budget']['replacement_error_scale']:.3f} < {loop_summary['cumulant_budget']['replacement_error_budget']:.2f}"],
        ["Derivative exponent", f"{loop_summary['cumulant_budget']['derivative_error_exponent']:.3f} < 0"],
        ["Runtime permission", loop_summary["permission"].replace("_", " ").lower()],
    ]
    add_table(doc, ["Audit item", "Value"], loop_rows, [3600, 5760])
    add_picture(
        doc,
        PROJECT / "figures" / "loop_equation_runtime_stability.png",
        6.5,
        "Figure S3. Loop-equation/Gronwall stability probe: residuals are allowed only when representation-transfer error remains within runtime budget.",
    )

    doc.add_heading("Branch-Factor / Volterra Path Probe", level=1)
    add_para(
        doc,
        "The branch-factor excerpt adds a branch-cut and collision-diagonal safety rule. A transformed artifact should remain in a declared half-plane sector, stay separated from other branch variables, keep a positive edge-phase gap, and make the Volterra fixed-point map contractive. The BBGKY appendix supplies the contact warning: pointwise identities away from collisions do not automatically govern near-collision cases.",
    )
    branch_rows = [
        ["Separated branch R0", f"{branch_summary['stable_r0']:.2f}"],
        ["Separated min separation/lambda", f"{branch_summary['stable_min_separation_ratio']:.3f}"],
        ["Separated min phase margin", f"{branch_summary['stable_min_scaled_phase_margin']:.3f}"],
        ["Separated Volterra ratio", f"{branch_summary['stable_volterra_contraction_ratio']:.4f} < 1"],
        ["Near-collision R0", f"{branch_summary['near_collision_r0']:.3f}"],
        ["Near-collision passes", str(branch_summary["near_collision_branch_passes"])],
        ["Runtime permission", branch_summary["permission"].replace("_", " ").lower()],
    ]
    add_table(doc, ["Audit item", "Value"], branch_rows, [3900, 5460])
    add_picture(
        doc,
        PROJECT / "figures" / "branch_factor_path_stability.png",
        6.5,
        "Figure S4. Branch-factor path stability probe: separated branches can be audited; near-collision source clusters are blocked for branch/transport review.",
    )

    doc.add_heading("Experimental Mathematics / Proof-Status Probe", level=1)
    add_para(
        doc,
        "The MathOverflow/PSLQ/Euler-sum discussion is used only as a proof-status analogy. Experimental mathematics can discover striking candidate identities, but a numerical relation, a remembered source, and a proof are different evidence objects. The medical LLM version of the same distinction is that literature-mined relations, plausible mechanisms, and fluent syntheses must remain labeled as hypotheses or surrogate evidence until validation status is established.",
    )
    add_para(
        doc,
        "The PhilPapers/PhilArchive NIETTU record is used in the same bounded way. A real archived or forthcoming topological unified-field-theory paper may be interesting as a source object, but it is not clinical validation for a medical LLM, multimodal documentation system, genetic-privacy policy, diagnosis, treatment, or patient-outcome claim. In the stress harness, that distinction becomes a provenance/proof-status gate: verified source existence is not the same thing as verified medical evidentiary currency.",
    )

    doc.add_heading("Multilingual Transcript Provenance Probe", level=1)
    add_para(
        doc,
        "The cross-script stress case treats raw transcript text as an evidence artifact that must pass provenance checks before clinical use. Mongolian Cyrillic, Hebrew right-to-left text, translated dubs, captions, OCR output, and encoding-damaged scrapes can all alter the information object before the LLM reasons over it. The runtime action is therefore audit-only until language, source, encoding, speaker, translation status, consent boundary, and clinical relevance are established.",
    )

    doc.add_heading("Survivorship-Bias Evaluation Probe", level=1)
    add_para(
        doc,
        "The survivorship-bias aircraft figure is included as a missing-denominator analogy, not as medical evidence. Medical LLM evaluation should not only count fluent answers that look plausible. It should also preserve prompts that fail, abstain, lose citations, drop subgroups, mistranslate source material, or never reach a terminal patient-outcome endpoint. Those invisible non-returning cases are often where the assurance boundary should be reinforced.",
    )

    doc.add_page_break()
    doc.add_heading("Reverse-Sprinkler Momentum-Flux Probe", level=1)
    add_para(
        doc,
        "The reverse-sprinkler result is included as an input-output asymmetry analogy, not as a biomedical mechanism. In the sprinkler problem, reversing the water flow does not simply reverse the ordinary mechanism; geometry controls the momentum flux that crosses the device. In medical LLM evaluation, the analogous warning is that answer-to-evidence reconstruction is not the inverse of evidence-to-answer generation. The runtime monitor should therefore preserve provenance, endpoint category, and boundary-crossing history rather than trusting a fluent retrospective explanation.",
    )

    doc.add_heading("Boundary Mathematics Note: HJB, Shocks, and KdV", level=1)
    add_para(
        doc,
        "A useful advanced analogy is to view runtime assurance as a control problem only while the value function remains well behaved. In that regime, a Hamilton-Jacobi-Bellman framing is natural: the monitor chooses an action that minimizes risk under the current evidence state. At claim-permission boundaries, however, the value surface can become nonsmooth. The correct mathematical language then looks closer to viscosity solutions, shocks, or distributional boundary conditions than to a smooth optimizer.",
    )
    add_para(
        doc,
        "KdV-style dispersive regularization is included only as an explanatory analogy. It suggests that a sharp boundary may be replaced by an oscillatory transition layer. In the capstone, that layer is the region where a surrogate endpoint, caption transcript, device metric, or workflow statistic begins to travel as if it were a hard patient outcome. The runtime monitor should detect and log that boundary-layer behavior before allowing the LLM to speak at a stronger claim level.",
    )

    doc.add_heading("Audiology/CODA Domain Probe", level=1)
    add_para(
        doc,
        "Audiology is a good future probe because every layer is a surrogate transformation: genetic etiology is not treatment preference, microphone capture is not the communicative act, captions are not the full encounter, speech perception is not quality of life, and cochlear-implant device performance is not autonomy. SIK3 is included only as a hearing-associated genetic hypothesis case; tinnitus GWAS evidence supports polygenic and neuropsychiatric links, not a tolerance or adaptation claim. CODA supplies the cultural reminder that Deaf communication and family life cannot be reduced to normalization-to-hearing metrics. This domain would force the runtime monitor to separate medical evidence, communication access, family authority, and patient-centered outcomes.",
    )
    add_para(
        doc,
        "Saliva evolutionary genomics adds a second genetics probe outside audiology. Pajic, Landau, Gokcumen, and Ruhl show that saliva-related SCPP genes have undergone duplication, loss, regulatory change, and signatures of selection in primate evolution. That is meaningful mechanism and baseline evidence for oral biology, diet, and host-microbe interaction. The runtime gate blocks the stronger move from evolutionary-genomics evidence to individual diagnosis, oral-disease prediction, or patient-outcome benefit unless separate clinical validation exists.",
    )

    doc.add_heading("Consumer Heart-Rate / Wearable Physiology Probe", level=1)
    add_para(
        doc,
        "The Harvard Health heart-rate source makes the wearable case concrete. Resting heart rate and exercise target zones can be useful longitudinal signals, but they are modified by age, fitness, stress, anxiety, hormones, medication, and activity level. The runtime monitor therefore narrows a wearable heart-rate claim to monitoring and clinician-follow-up language. It blocks diagnosis, reassurance, and individualized exercise prescription unless separate clinical context and validation are available.",
    )

    doc.add_heading("Altitude, Circadian, and Dinacharya Probe", level=1)
    add_para(
        doc,
        "Altitude and circadian rhythm are represented as transport modifiers. A hearing, tinnitus, or sleep model that is plausible at sea level may need stress testing under hypoxia, hypobaria, altered sleep timing, or culturally structured daily routines such as Dinacharya. The monitor therefore treats altitude and daily-rhythm inputs as context features that can trigger generalization testing.",
    )

    doc.add_heading("Robot-as-Cadaver Probe", level=1)
    add_para(
        doc,
        "The robotic-simulation case treats robots, synthetic organs, and humanoid patient simulators as cadaver-like rehearsal surfaces: useful for repeatable practice and audit, but not proof that the system improves patient outcomes. The Better Than Us reference is used only as a scope-setting cultural analogy for embodied robots in medicine; the evidence gate relies on robotic surgery simulation and humanoid patient-robot literature.",
    )

    doc.add_heading("Neuropsychology Norm-Drift Probe", level=1)
    add_para(
        doc,
        "The reverse-Flynn case tests whether the monitor notices temporal and population drift in psychometric interpretation. WAIS scores in a high-secure psychiatric setting cannot be interpreted as ordinary general-population evidence without checking decade, test version, processing-speed profile, repeat-testing reliability, and clinical pathway implications.",
    )

    doc.save(path)
    return path


def write_readme(
    summary: dict,
    zdd_summary: dict,
    model_summary: dict,
    depletion_summary: dict,
    projection_summary: dict,
    kelly_summary: dict,
    witness_summary: dict,
    consent_summary: dict,
    event_summary: dict,
    measure_summary: dict,
    mahalanobis_summary: dict,
    matrix_summary: dict,
    tail_summary: dict,
    hex_summary: dict,
    hex_scaling_summary: dict,
    proof_status_summary: dict,
    picture_summary: dict,
    loop_summary: dict,
    branch_summary: dict,
    docx_paths: list[Path],
    pdf_paths: list[Path],
) -> Path:
    path = PACKAGE / "README.md"
    lines = [
        "# Module 14 Capstone Package",
        "",
        "## Project",
        "",
        "Surrogate-Aware Runtime Assurance for Agentic Medical LLMs: Preventing evidence overclaiming in clinician-facing AI systems.",
        "",
        "## Goal",
        "",
        "This package proposes and implements a synthetic runtime-assurance framework for agentic medical LLM outputs and proposed actions, with the practical genetics scope narrowed to recurring hospital problems: pharmacogenomic alerts, newborn-screening follow-up, ACMG secondary findings, hereditary cancer/FH/Lynch flags, VUS handling, and family-linked privacy boundaries. The central idea is to decouple evidence assets from claim currencies: a fluent answer, real citation, surrogate endpoint, workflow metric, genetic result, and patient-outcome RCT do not carry the same permission level.",
        "",
        "## What Is Already Done",
        "",
        f"- Stress-test harness implemented with {summary['scenario_count']} scenarios.",
        f"- Expected actions matched in {summary['matched_expected_count']} of {summary['scenario_count']} scenarios.",
        f"- ZDD-style sparse-family audit implemented across {zdd_summary['feature_universe_size']} evidence features.",
        f"- Observed stress states compressed to {zdd_summary['observed_zdd_node_count']} ZDD nodes versus a naive trie upper bound of {zdd_summary['naive_observed_trie_upper_bound']}.",
        f"- Model-selection audit implemented with {model_summary['candidate_model_count']} candidate claim models and {model_summary['matched_expected_count']} matched expected actions.",
        f"- Proof-status/provenance poset audit enumerates {proof_status_summary['states_enumerated']} states and {proof_status_summary['cover_transitions_enumerated']} cover transitions; only {proof_status_summary['permission_counts']['hard_outcome_allowed_with_caveats']} states allow hard outcomes, and NIETTU remains audit-only under source/proof upgrades.",
        f"- Picture-language diagram audit classifies {picture_summary['diagram_count']} multimodal/mathematical artifacts; none allow hard-outcome permission, {picture_summary['transfer_reset_diagram_count']} reset validation after representation transfer, and the Bourgade-Huang loop-equation row remains universality-audit-only.",
        f"- Loop-equation/Gronwall stability probe checks {loop_summary['required_gate_count']} invariant gates: stable transfer ends at {loop_summary['stable_final_error']:.4f} under budget {loop_summary['stable_budget']:.3f}, while missing-provenance transfer ends at {loop_summary['unstable_final_error']:.4f} over budget {loop_summary['unstable_budget']:.3f}; switching-cumulant checks cancel the main quadratic term and keep rare-event/replacement errors inside budget.",
        f"- Branch-factor/Volterra path probe accepts the separated synthetic branch with contraction ratio {branch_summary['stable_volterra_contraction_ratio']:.4f} and blocks the near-collision family/population branch, preserving the BBGKY contact/collision warning.",
        f"- Sequential-depletion ordering experiment implemented across {depletion_summary['permutation_count']} runtime-load permutations.",
        f"- Orthogonal-projection confounding experiment passed: max design inner product after projection is {projection_summary['max_abs_design_inner_product_after']:.2e}.",
        f"- Kelly-style runtime exposure experiment shows all-in ruin probability {kelly_summary['all_in_ruin_probability']:.3f} despite favorable expected value.",
        f"- Uniform-witness sampling experiment shows biased strongest-five witness display has variation distance {witness_summary['variation_distance_from_uniform']:.3f} from the uniform admissible witness distribution.",
        f"- Family genomic-consent experiment detects ranked-choice cycling, blocks full raw release, and preserves partial inclusion under synthetic epsilon budget {consent_summary['runtime_budget']['epsilon']:.2f}.",
        f"- Finite event-algebra experiment verifies measurable-event identities over {event_summary['sample_space_size']} runtime histories.",
        f"- Measure-on-measures experiment compares {measure_summary['source_measure_count']} source-population measures against a meta-measure mixture.",
        f"- Mahalanobis covariate-distance experiment flags {len(mahalanobis_summary['flagged_sources'])} synthetic source(s) for source-specific validation.",
        f"- Exact matrix counterexample probe verifies a 2x2 rank-one projector failure with violation factor {matrix_summary['violation_factor_exact']}; included only as a counterexample-search analogy.",
        f"- Tail/maximal-inequality runtime probe enumerates {tail_summary['path_count']} mean-zero paths and shows pathwise boundary crossing can exceed terminal-only exceedance.",
        f"- Hex/Y boundary probe enumerates {hex_summary['total_full_boards_enumerated']} full boards with no both/neither terminal crossings and verifies majority-triangle coarse-graining has {hex_summary['majority_triangle_summary']['tie_count']} local ties.",
        f"- Hex scaling/coarse-graining probe samples {hex_scaling_summary['sampled_full_boards']} larger boards with {hex_scaling_summary['ambiguous_terminal_count']} ambiguous terminals, then shows generic smoothing can flip the global crossing at rate {hex_scaling_summary['max_coarse_grain_flip_rate']:.3f}.",
        "- Companion Python verification script included for sequential-depletion manuscript checks.",
        "- Figures generated for gate architecture, runtime action counts, endpoint/design permission mapping, proof-status poset audit, and supporting experiment probes.",
        "- Survivorship-bias aircraft figure added as a missing-denominator analogy for medical LLM evaluation, with original SVG and PNG included under its CC BY-SA 4.0 attribution.",
        "- Reverse-sprinkler momentum-flux result added as a bounded analogy for why LLM evidence flow is not safely invertible without provenance logs.",
        "- Tests verify that surrogate evidence cannot become a hard-outcome claim, fabricated citations are denied, CLOT-style low acceptance is routed to workflow diagnosis, Kanazawa-style evidence ambiguity is narrowed before claim composition, recurring hospital genetics triage, SIK3/tinnitus, saliva evolutionary genomics, and consumer heart-rate/wearable physiology are narrowed to surrogate monitoring or mechanism/association evidence, altitude/circadian and reverse-Flynn cases are stress-tested, robot-as-cadaver rehearsal stays training-only, cross-script, Romani-style language, and picture-language inputs require provenance audit, and Connes/Penrose evidence-patch collapse is blocked by a transport gate.",
        "- Liquid neural networks noted as a future streaming-monitor architecture while keeping the assurance layer model-agnostic.",
        "- Nonlocal modeling, game p-Laplacian inpainting, and coding/cryptography references added as bounded analogies for memory, reconstruction, and decodable provenance in transformed medical inputs.",
        "- PSLQ/Euler-sum experimental mathematics added as a bounded analogy for candidate discovery versus proof/validation status.",
        "- PhilPapers/PhilArchive NIETTU topological-theory record added as a proof-status/provenance stress case: verified source existence does not imply clinical validation or medical claim permission.",
        "- Kolmogorov-Arnold Networks critical-assessment preprint added as a model-selection stress case: theorem-inspired architecture claims remain method/surrogate evidence until clinical validation gates are satisfied.",
        "- Jaffe-Liu picture-language program added as a multimodal simulation stress case: diagrams, screenshots, photos, and videos require an explicit simulation map before clinical claim use.",
        "- Bourgade-Huang loop-equation characterization added as a universality stress case: approximate invariant hierarchies support portability audits but do not authorize clinical outcome claims.",
        "- Gronwall, resolvent-stability, cumulant-error, and random d-regular switching-calculus excerpts added as runtime-stability language for error propagation and sparse-network perturbation audits.",
        "- Branch-factor, radial Volterra, and BBGKY collision excerpts added as branch-stability language for near-collision family, population, and contact-term audits.",
        "- Axelrod/Hamilton and Wu/Axelrod cooperation/noise references added as bounded adoption and family-consent analogies; noisy interaction is audited before being interpreted as defection/refusal.",
        "- Sequential depletion ordering added as a concrete runtime-budget experiment: small-first, large-first, FIFO, and LIFO are not interchangeable once residual budgets matter.",
        "- Figaro noted as a future probabilistic-programming substrate for richer runtime-history and source-measure models; the submitted artifact remains dependency-light Python.",
        "",
        "## Package Contents",
        "",
        "- `proposal/`: Week 11 research proposal in DOCX and PDF.",
        "- `paper/`: Final paper draft in DOCX and PDF.",
        "- `summary/`: Required 1-2 page summary sheet in DOCX and PDF.",
        "- `supplement/`: Technical supplement with scenario matrix and reproducibility notes.",
        "- `code/`: Synthetic stress-test harness and tests.",
        "- `results/`: JSON and CSV outputs, including the ZDD sparse-family and model-selection summaries.",
        "- `figures/`: Generated figures used in the paper.",
        "- `demo/`: Skater/Hawkes logged replay video and contact sheet, included only as a runtime-assurance visual analogy.",
        "- Runtime implementation note: async/coroutine control flow is used as an analogy for suspending output until assurance checks pass.",
        "",
        "## Safety Boundary",
        "",
        "This is not a medical product. It does not use real patient data, does not diagnose, does not recommend treatment, and does not validate clinical effectiveness. It is an auditable design and evaluation artifact for clinician-facing LLM safety.",
        "",
        "## Reproduction",
        "",
        "```powershell",
        "$env:PYTHONPATH=(Resolve-Path 'work\\evidence_gated_llm_capstone').Path",
        "& 'C:\\Users\\anaxe\\.cache\\codex-runtimes\\codex-primary-runtime\\dependencies\\python\\python.exe' 'work\\evidence_gated_llm_capstone\\run_evidence_gate_stress_test.py'",
        "& 'C:\\Users\\anaxe\\.cache\\codex-runtimes\\codex-primary-runtime\\dependencies\\python\\python.exe' 'work\\evidence_gated_llm_capstone\\zdd_sparse_claim_family.py'",
        "& 'C:\\Users\\anaxe\\.cache\\codex-runtimes\\codex-primary-runtime\\dependencies\\python\\python.exe' 'work\\evidence_gated_llm_capstone\\model_selection_claim_policy.py'",
        "& 'C:\\Users\\anaxe\\.cache\\codex-runtimes\\codex-primary-runtime\\dependencies\\python\\python.exe' 'work\\evidence_gated_llm_capstone\\proof_status_poset_experiment.py'",
        "& 'C:\\Users\\anaxe\\.cache\\codex-runtimes\\codex-primary-runtime\\dependencies\\python\\python.exe' 'work\\evidence_gated_llm_capstone\\picture_language_diagram_audit.py'",
        "& 'C:\\Users\\anaxe\\.cache\\codex-runtimes\\codex-primary-runtime\\dependencies\\python\\python.exe' 'work\\evidence_gated_llm_capstone\\loop_equation_runtime_stability.py'",
        "& 'C:\\Users\\anaxe\\.cache\\codex-runtimes\\codex-primary-runtime\\dependencies\\python\\python.exe' 'work\\evidence_gated_llm_capstone\\branch_factor_path_stability.py'",
        "& 'C:\\Users\\anaxe\\.cache\\codex-runtimes\\codex-primary-runtime\\dependencies\\python\\python.exe' 'work\\evidence_gated_llm_capstone\\depletion_ordering_experiment.py'",
        "& 'C:\\Users\\anaxe\\.cache\\codex-runtimes\\codex-primary-runtime\\dependencies\\python\\python.exe' 'work\\evidence_gated_llm_capstone\\orthogonal_projection_experiment.py'",
        "& 'C:\\Users\\anaxe\\.cache\\codex-runtimes\\codex-primary-runtime\\dependencies\\python\\python.exe' 'work\\evidence_gated_llm_capstone\\kelly_runtime_budget_experiment.py'",
        "& 'C:\\Users\\anaxe\\.cache\\codex-runtimes\\codex-primary-runtime\\dependencies\\python\\python.exe' 'work\\evidence_gated_llm_capstone\\uniform_witness_sampling_experiment.py'",
        "& 'C:\\Users\\anaxe\\.cache\\codex-runtimes\\codex-primary-runtime\\dependencies\\python\\python.exe' 'work\\evidence_gated_llm_capstone\\consent_aggregation_experiment.py'",
        "& 'C:\\Users\\anaxe\\.cache\\codex-runtimes\\codex-primary-runtime\\dependencies\\python\\python.exe' 'work\\evidence_gated_llm_capstone\\event_algebra_experiment.py'",
        "& 'C:\\Users\\anaxe\\.cache\\codex-runtimes\\codex-primary-runtime\\dependencies\\python\\python.exe' 'work\\evidence_gated_llm_capstone\\measure_on_measures_experiment.py'",
        "& 'C:\\Users\\anaxe\\.cache\\codex-runtimes\\codex-primary-runtime\\dependencies\\python\\python.exe' 'work\\evidence_gated_llm_capstone\\mahalanobis_covariate_experiment.py'",
        "& 'C:\\Users\\anaxe\\.cache\\codex-runtimes\\codex-primary-runtime\\dependencies\\python\\python.exe' 'work\\evidence_gated_llm_capstone\\noncommutative_amgm_counterexample.py'",
        "& 'C:\\Users\\anaxe\\.cache\\codex-runtimes\\codex-primary-runtime\\dependencies\\python\\python.exe' 'work\\evidence_gated_llm_capstone\\tail_maximal_inequality_experiment.py'",
        "& 'C:\\Users\\anaxe\\.cache\\codex-runtimes\\codex-primary-runtime\\dependencies\\python\\python.exe' 'work\\evidence_gated_llm_capstone\\hex_boundary_invariant_experiment.py'",
        "& 'C:\\Users\\anaxe\\.cache\\codex-runtimes\\codex-primary-runtime\\dependencies\\python\\python.exe' 'work\\evidence_gated_llm_capstone\\hex_scaling_coarse_grain_experiment.py'",
        "& 'C:\\Users\\anaxe\\.cache\\codex-runtimes\\codex-primary-runtime\\dependencies\\python\\python.exe' 'work\\evidence_gated_llm_capstone\\sequential_depletion_verification.py'",
        "py -3.13 -m pytest 'work\\evidence_gated_llm_capstone' -q",
        "```",
        "",
        "## Final DOCX Files",
        "",
    ]
    for p in docx_paths:
        lines.append(f"- `{p.name}`")
    lines.extend(["", "## Final PDF Files", ""])
    for p in pdf_paths:
        lines.append(f"- `{p.name}`")
    path.write_text("\n".join(lines) + "\n", encoding="utf-8")
    return path


def copy_assets() -> None:
    for name in ["code", "results", "figures", "demo"]:
        (PACKAGE / name).mkdir(parents=True, exist_ok=True)
    for file in [
        "run_evidence_gate_stress_test.py",
        "test_evidence_gate_stress_test.py",
        "zdd_sparse_claim_family.py",
        "test_zdd_sparse_claim_family.py",
        "model_selection_claim_policy.py",
        "test_model_selection_claim_policy.py",
        "proof_status_poset_experiment.py",
        "test_proof_status_poset_experiment.py",
        "picture_language_diagram_audit.py",
        "test_picture_language_diagram_audit.py",
        "loop_equation_runtime_stability.py",
        "test_loop_equation_runtime_stability.py",
        "branch_factor_path_stability.py",
        "test_branch_factor_path_stability.py",
        "depletion_ordering_experiment.py",
        "test_depletion_ordering_experiment.py",
        "orthogonal_projection_experiment.py",
        "test_orthogonal_projection_experiment.py",
        "kelly_runtime_budget_experiment.py",
        "test_kelly_runtime_budget_experiment.py",
        "uniform_witness_sampling_experiment.py",
        "test_uniform_witness_sampling_experiment.py",
        "consent_aggregation_experiment.py",
        "test_consent_aggregation_experiment.py",
        "event_algebra_experiment.py",
        "test_event_algebra_experiment.py",
        "measure_on_measures_experiment.py",
        "test_measure_on_measures_experiment.py",
        "mahalanobis_covariate_experiment.py",
        "test_mahalanobis_covariate_experiment.py",
        "noncommutative_amgm_counterexample.py",
        "test_noncommutative_amgm_counterexample.py",
        "tail_maximal_inequality_experiment.py",
        "test_tail_maximal_inequality_experiment.py",
        "hex_boundary_invariant_experiment.py",
        "test_hex_boundary_invariant_experiment.py",
        "hex_scaling_coarse_grain_experiment.py",
        "test_hex_scaling_coarse_grain_experiment.py",
        "sequential_depletion_verification.py",
    ]:
        shutil.copy2(PROJECT / file, PACKAGE / "code" / file)
    for file in (PROJECT / "results").glob("*"):
        shutil.copy2(file, PACKAGE / "results" / file.name)
    for file in (PROJECT / "figures").glob("*"):
        shutil.copy2(file, PACKAGE / "figures" / file.name)

    video = Path(r"C:\Users\anaxe\Downloads\skater_hawkes_logged_replay.mp4")
    if video.exists():
        shutil.copy2(video, PACKAGE / "demo" / video.name)
    contact_sheet = WORK / "video_inspect" / "skater_hawkes_logged_replay" / "contact_sheet.png"
    if contact_sheet.exists():
        shutil.copy2(contact_sheet, PACKAGE / "demo" / "skater_hawkes_contact_sheet.png")


def write_manifest() -> Path:
    files = [p for p in PACKAGE.rglob("*") if p.is_file() and p.name != "MANIFEST_SHA256.csv"]
    path = PACKAGE / "MANIFEST_SHA256.csv"
    with path.open("w", newline="", encoding="utf-8") as f:
        writer = csv.writer(f)
        writer.writerow(["path", "sha256", "bytes"])
        for file in sorted(files):
            writer.writerow([str(file.relative_to(PACKAGE)).replace("\\", "/"), sha256(file), file.stat().st_size])
    return path


def zip_package() -> Path:
    zip_path = OUTPUTS / f"{PACKAGE.name}.zip"
    if zip_path.exists():
        zip_path.unlink()
    with zipfile.ZipFile(zip_path, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        for file in sorted(PACKAGE.rglob("*")):
            if file.is_file():
                zf.write(file, file.relative_to(PACKAGE.parent))
    return zip_path


def main() -> None:
    if PACKAGE.exists():
        shutil.rmtree(PACKAGE)
    PACKAGE.mkdir(parents=True, exist_ok=True)

    summary = run_stress_test()
    zdd_summary = run_zdd_audit()
    model_summary = run_model_selection_audit()
    depletion_summary = run_depletion_ordering_experiment()
    projection_summary = run_orthogonal_projection_experiment()
    kelly_summary = run_kelly_runtime_budget_experiment()
    witness_summary = run_uniform_witness_sampling_experiment()
    consent_summary = run_consent_aggregation_experiment()
    event_summary = run_event_algebra_experiment()
    measure_summary = run_measure_on_measures_experiment()
    mahalanobis_summary = run_mahalanobis_covariate_experiment()
    matrix_summary = run_noncommutative_amgm_counterexample()
    tail_summary = run_tail_maximal_inequality_experiment()
    hex_summary = run_hex_boundary_invariant_experiment()
    hex_scaling_summary = run_hex_scaling_coarse_grain_experiment()
    proof_status_summary = run_proof_status_poset_experiment()
    picture_summary = run_picture_language_diagram_audit()
    loop_summary = run_loop_equation_runtime_stability()
    branch_summary = run_branch_factor_path_stability()
    proposal = create_proposal(
        summary,
        zdd_summary,
        model_summary,
        depletion_summary,
        projection_summary,
        kelly_summary,
        witness_summary,
        consent_summary,
        event_summary,
        measure_summary,
        mahalanobis_summary,
        hex_scaling_summary,
        proof_status_summary,
        picture_summary,
        loop_summary,
        branch_summary,
    )
    paper = create_final_paper(
        summary,
        zdd_summary,
        model_summary,
        depletion_summary,
        projection_summary,
        kelly_summary,
        witness_summary,
        consent_summary,
        event_summary,
        measure_summary,
        mahalanobis_summary,
        matrix_summary,
        tail_summary,
        hex_summary,
        hex_scaling_summary,
        proof_status_summary,
        picture_summary,
        loop_summary,
        branch_summary,
    )
    sheet = create_summary_sheet(
        summary,
        zdd_summary,
        model_summary,
        depletion_summary,
        projection_summary,
        kelly_summary,
        witness_summary,
        consent_summary,
        event_summary,
        measure_summary,
        mahalanobis_summary,
    )
    supplement = create_technical_supplement(
        summary,
        zdd_summary,
        model_summary,
        depletion_summary,
        projection_summary,
        kelly_summary,
        witness_summary,
        consent_summary,
        event_summary,
        measure_summary,
        mahalanobis_summary,
        matrix_summary,
        tail_summary,
        hex_summary,
        hex_scaling_summary,
        proof_status_summary,
        picture_summary,
        loop_summary,
        branch_summary,
    )
    docx_paths = [proposal, paper, sheet, supplement]
    pdf_paths = []
    for docx_path in docx_paths:
        export_pdf(docx_path)
        pdf_paths.append(docx_path.with_suffix(".pdf"))

    copy_assets()
    write_readme(
        summary,
        zdd_summary,
        model_summary,
        depletion_summary,
        projection_summary,
        kelly_summary,
        witness_summary,
        consent_summary,
        event_summary,
        measure_summary,
        mahalanobis_summary,
        matrix_summary,
        tail_summary,
        hex_summary,
        hex_scaling_summary,
        proof_status_summary,
        picture_summary,
        loop_summary,
        branch_summary,
        docx_paths,
        pdf_paths,
    )
    write_manifest()
    zip_path = zip_package()

    print(json.dumps({
        "package_dir": str(PACKAGE),
        "zip": str(zip_path),
        "docx": [str(p) for p in docx_paths],
        "pdf": [str(p) for p in pdf_paths],
        "scenario_count": summary["scenario_count"],
        "matched_expected_count": summary["matched_expected_count"],
        "zdd_node_count": zdd_summary["observed_zdd_node_count"],
        "model_selection_candidates": model_summary["candidate_model_count"],
        "model_selection_matched_expected": model_summary["matched_expected_count"],
        "depletion_permutation_count": depletion_summary["permutation_count"],
        "depletion_context_reversal_observed": depletion_summary["context_reversal_observed"],
        "projection_orthogonality_passed": projection_summary["orthogonality_passed"],
        "kelly_all_in_ruin_probability": kelly_summary["all_in_ruin_probability"],
        "uniform_witness_variation_distance": witness_summary["variation_distance_from_uniform"],
        "consent_cycle_detected": consent_summary["condorcet_cycle_detected"],
        "partial_inclusion_runtime_safe": consent_summary["partial_inclusion_runtime_safe"],
        "event_union_identity_error": event_summary["union_identity_error"],
        "measure_on_measures_max_tv": measure_summary["max_source_tv_to_mixture"],
        "mahalanobis_flagged_sources": mahalanobis_summary["flagged_sources"],
        "matrix_counterexample_violation_factor": matrix_summary["violation_factor_exact"],
        "matrix_counterexample_pd_margin": matrix_summary["positive_definite_margin_exact"],
        "tail_maximal_crossing_probability": tail_summary["exact_crossing_probability"],
        "tail_maximal_bound": tail_summary["kolmogorov_bound"],
        "hex_full_boards_enumerated": hex_summary["total_full_boards_enumerated"],
        "hex_no_both_or_neither": hex_summary["exactly_one_crossing_all_full_boards"],
        "hex_scaling_sampled_full_boards": hex_scaling_summary["sampled_full_boards"],
        "hex_scaling_ambiguous_terminals": hex_scaling_summary["ambiguous_terminal_count"],
        "hex_scaling_max_smoothing_flip_rate": hex_scaling_summary["max_coarse_grain_flip_rate"],
        "proof_status_states_enumerated": proof_status_summary["states_enumerated"],
        "proof_status_hard_outcome_states": proof_status_summary["permission_counts"]["hard_outcome_allowed_with_caveats"],
        "proof_status_niettu_permission": proof_status_summary["niettu_permission"],
        "picture_language_diagram_count": picture_summary["diagram_count"],
        "picture_language_hard_outcome_permissions": picture_summary["hard_outcome_permission_count"],
        "loop_stable_chain_passes": loop_summary["stable_chain_passes"],
        "loop_unstable_chain_passes": loop_summary["unstable_chain_passes"],
        "loop_main_cancellation_passes": loop_summary["cumulant_budget"]["main_cancellation_passes"],
        "loop_bad_event_negligible": loop_summary["cumulant_budget"]["bad_event_negligible"],
        "branch_stable_path_passes": branch_summary["stable_branch_passes"],
        "branch_near_collision_path_passes": branch_summary["near_collision_branch_passes"],
    }, indent=2))


if __name__ == "__main__":
    main()
